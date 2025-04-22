import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QMessageBox
from PyQt5.uic import loadUi
import os
from PyQt5.QtCore import QStringListModel
from PyQt5.QtWidgets import QListWidgetItem
from db.database import init_database
from ui.main_window_ui import Ui_MainWindow
import json
from PyQt5.QtGui import QIcon
from PyQt5.QtGui import QBrush, QColor
from export_fiches import export_fiches
from PyQt5 import QtGui
from PyQt5.QtCore import Qt
from PyQt5 import QtWidgets
import os
import json
from PyQt5.QtWidgets import QCompleter
from docx import Document
from PyQt5.QtWidgets import QMessageBox
import sqlite3
from PyQt5.QtWidgets import QHeaderView
from PyQt5.QtWidgets import QMenu
from db import crud
import xlsxwriter 
from controllers.job_controller import JobController
from controllers.competence_controller import CompetenceController
from controllers.referentiel_controller import ReferentielController
from controllers.import_export_controller import ImportExportController
from controllers.openai_controller import OpenAIController
from controllers.fichePoste_controller import FichePosteController
from controllers.FinaliteEmploiController import FinaliteEmploiController
from PyQt5.QtWidgets import QInputDialog
from db import crud
from ui.dialogs.dialog_parametres_emploi import DialogParametresEmploi
from PyQt5.QtWidgets import QTreeWidgetItem
from PyQt5 import QtCore
import re
from ui.dialogs.dialog_nouveau_projet import DialogNouveauProjet
from PyQt5.QtWidgets import QDialog
import os
import pandas as pd
from docx import Document
from PyQt5.QtWidgets import QFileDialog
import pandas as pd
from PyQt5.QtWidgets import QFileDialog, QMessageBox, QTableWidgetItem





class MainApp(QMainWindow):
    def __init__(self):
        super(MainApp, self).__init__()

        # Chemin vers le fichier .ui
        ui_path = os.path.join("ui", "main_window.ui")
        self.historique_carto = []  # Conserve l'historique des échanges liés à la cartographie
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        
        # Charger et appliquer le style
        with open("resources/styles/style.qss", "r") as file:
            style = file.read()
            app.setStyleSheet(style)
        
        # Contrôleurs
        # 🔸 Pas de projet actif au démarrage
        self.customize_layouts()
        self.current_projet_id = None
        self.job_controller = None
        self.finalite_controller = None 
        self.fiche_poste_controller=None
        self.competence_controller = None
        self.openai_controller = None
        self.ui.stackedWidget.setCurrentIndex(0)
        self.contenu_fichier_client = ""
        self.type_fichier_client = ""
        self.chemin_fichier_unique = None
        self.description_fichier_unique = None
        self.interruption_generation = False
        ####self.ui.btn_importer_cartographie_competences_emplois.setEnabled(False)

        # 🔸 Désactiver les pages tant que pas de projet
        self._verrouiller_pages()
        self.ui.btn_ajouter_zone_fichier.setEnabled(True)
        self.ui.combo_type_competence.addItems(["hard", "soft"])
        self.ui.combo_type_fichier.addItems([
            "Offre d'emploi",
            "Fiche de poste",
            "Fiche métier",
            "Autre"
        ])
        tree = self.ui.tree_cartographie_metiers
        tree.setDragDropMode(tree.InternalMove)
        tree.setSelectionMode(tree.ExtendedSelection)
        tree.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        tree.customContextMenuRequested.connect(self.show_tree_context_menu)
        tree.itemDoubleClicked.connect(self.edit_tree_item)
        tree.itemSelectionChanged.connect(self.mettre_a_jour_etat_boutons_tree)
        self.ui.table_competences.setColumnCount(3)
        self.ui.table_competences.setHorizontalHeaderLabels(["Compétence", "Niveau requis", "Type"])
        self.ui.table_competences_fichedeposte.setColumnCount(3)
        self.ui.table_competences_fichedeposte.setHorizontalHeaderLabels(["Compétence", "Niveau requis", "Type"])
        self.ui.table_dictionnaire_competences.horizontalHeader().setStretchLastSection(True)
        self.ui.table_dictionnaire_competences.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.ui.table_dictionnaire_competences.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.ui.table_niveaux_competence.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.ui.table_competences.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.ui.table_competences_fichedeposte.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.ui.tree_carto_emploi_fichedeposte.setHeaderLabel("Veuillez sélectionner un emploi : ")
        self.ui.btn_exporter_carto_competence_emploi.clicked.connect(self.exporter_carto_competences_emploi)
        self.ui.btn_lancer_generateur_activites.setEnabled(False)

        self.model_competences = QStringListModel()
        self.completer = QCompleter(self.model_competences, self)
        self.completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.ui.input_competence.setCompleter(self.completer)
        # 🔸 Connecter boutons
        self.ui.table_competences.itemSelectionChanged.connect(self.on_selection_change)
        self.ui.btn_sauvegarder_liste_competence.clicked.connect(self.sauvegarder_ligne_modifiee)
        self.ui.btn_supprimer_competence.clicked.connect(self.supprimer_competences_selectionnees_competencesemploi)
        self.ui.input_competence.textEdited.connect(self.mettre_a_jour_autocompletion_competence)
        self.ui.btn_ajouter_competence.clicked.connect(self.ajouter_competence_a_emploi)
        self.ui.btn_ajouter_zone_fichier.clicked.connect(self.ajouter_zone_fichier)
        self.ui.btn_modifier_inputs_projet.clicked.connect(lambda: self.set_inputs_projet_enabled(True))
        self.ui.btn_lancer_generateur_activites.clicked.connect(self.generer_fiche_de_poste_via_ia)
        self.ui.btn_sauvegarder_fiche_poste.clicked.connect(self.sauvegarder_fiche_de_poste)
        self.ui.combo_sous_famille.currentIndexChanged.connect(self.charger_competences_sous_famille)
        self.ui.tree_cartographie_metier_generateur_competence.itemClicked.connect(self.on_tree_generateur_item_clicked)
        self.ui.tree_carto_emploi_fichedeposte.itemClicked.connect(self.on_tree_fiche_poste_item_clicked)
        # Lors de l'initialisation
        self.ui.tree_cartographie_metier_generateur_competence.itemClicked.connect(self.on_emploi_selectionne)

        self.ui.tree_carto_emploi_fichedeposte.itemSelectionChanged.connect(self.on_selection_changed)
        self.ui.btn_exporter_fiches.clicked.connect(self.export_fiches)
        #self.btn_exporter_fiches.clicked.connect(self.export_fiches)
        self.ui.btn_referentiel_2.clicked.connect(self.update_combo_sous_famille)
        #self.ui.btn_ajouter_famille.clicked.connect(self.ajouter_famille)
        self.ui.btn_ajouter_familles.clicked.connect(self.ajouter_familles_par_liste)
        self.ui.btn_ajouter_sous_familles.clicked.connect(self.ajouter_sous_familles_par_liste)
        self.ui.btn_ajouter_metier.clicked.connect(self.ajouter_metiers_par_liste)
        self.ui.btn_exporter_cartographie.clicked.connect(self.exporter_cartographie_en_excel)
        self.ui.btn_nouveau_projet.clicked.connect(self.creer_nouveau_projet)
        self.ui.btn_ouvrir_projet.clicked.connect(self.ouvrir_projet)
        self.ui.btn_sauvegarder_inputs_projet.clicked.connect(self.sauvegarder_parametres_projet)
        self.ui.btn_importer_fichier_client.clicked.connect(self.importer_fichier_client)
        self.ui.btn_envoyer_ia_carto_emplois.clicked.connect(self.discuter_avec_ia_cartographie)
        self.ui.btn_sauvegarder_cartographie.clicked.connect(self.sauvegarder_cartographie_actuelle)
        self.ui.btn_cartographie.clicked.connect(self.on_cartographie_emploi_toggled)
        self.ui.btn_fiches_poste.clicked.connect(self.on_generateur_competences_emploi_toggled)
        self.ui.btn_referentiel.clicked.connect(self.on_referentiel_toggled)
        self.ui.btn_dictionnaire.clicked.connect(self.on_dictionnaire_toggled)
        self.ui.btn_referentiel_2.clicked.connect(self.on_carto_niveau_competence_toggled)
        self.ui.btn_dictionnaire_2.clicked.connect(self.on_fiche_poste_toggled)
        self.ui.btn_parametres.clicked.connect(self.on_parametres_toggled)
        self.ui.btn_importer_referentiel_competences.clicked.connect(self.importer_referentiel_competences_excel)
        self.ui.btn_generer_referentiel_competences_ia.clicked.connect(self.generer_referentiel_via_ia)
        self.ui.btn_ajouter_competence_referentiel.clicked.connect(self.ajouter_competence_referentiel)
        self.ui.btn_supprimer_competence_referentiel.clicked.connect(self.supprimer_competences_selectionnees)
        self.ui.btn_fusionner_competence_referentiel.clicked.connect(self.fusionner_competences_selectionnees)
        self.ui.btn_sauvegarder_referentiel_competence.clicked.connect(self.sauvegarder_referentiel_competence)
        self.ui.btn_exporter_referentiel_competence.clicked.connect(self.exporter_referentiel_competence_excel)
        self.ui.btn_importer_cartographie.clicked.connect(self.importer_cartographie_depuis_excel)

        # === Onglet Dictionnaire ===
        self.ui.btn_importer_dictionnaire_competences.clicked.connect(self.importer_dictionnaire_excel)
        self.ui.btn_generer_dictionnaire_ia_2.clicked.connect(self.generer_dictionnaire_via_ia)
        self.ui.btn_ajouter_comptence_dictionnaire.clicked.connect(self.ajouter_ligne_dictionnaire)
        self.ui.btn_supprimer_competence_disctionnaire.clicked.connect(self.supprimer_ligne_dictionnaire)
        self.ui.btn_sauvegarder_dictionnaire.clicked.connect(self.sauvegarder_dictionnaire)
        self.ui.btn_exporter_dictionnaire.clicked.connect(self.exporter_dictionnaire_excel)
        self.ui.btn_generer_dictionnaire_ia.clicked.connect(self.sauvegarder_definition_niveau)
        self.ui.combo_niveau_competence_a_definir.currentIndexChanged.connect(self.charger_definition_niveau_selectionne)
        
        self.ui.btn_generateur_competences.clicked.connect(self.lancer_generation_competences)
        self.chemin_dossier_fichiers = ""
        self.ui.btn_choisir_dossier_fichiers.clicked.connect(self.choisir_dossier_fichiers)
        self.ui.btn_generer_cartographie.clicked.connect(self.generer_cartographie_via_ia)

        self.setWindowTitle("Application Gestion des Compétences")
        self.setMinimumSize(1000, 700)
        self.show()

    def on_selection_change(self):
        has_selection = bool(self.ui.table_competences.selectedItems())
        self.ui.btn_sauvegarder_liste_competence.setEnabled(has_selection)
        self.ui.btn_supprimer_competence.setEnabled(has_selection)

    def sauvegarder_ligne_modifiee(self):
        row = self.ui.table_competences.currentRow()
        if row == -1:
            QMessageBox.warning(self, "Aucune sélection", "Veuillez sélectionner une ligne à enregistrer.")
            return

        nom = self.ui.table_competences.item(row, 0).text().strip()
        niveau = self.ui.table_competences.item(row, 1).text().strip()
        type_ = self.ui.table_competences.item(row, 2).text().strip()

        # Appel au contrôleur pour modifier la compétence
        try:
            self.competence_controller.modifier_competence_globale(self.selected_emploi, nom, type_, niveau)
            QMessageBox.information(self, "Succès", "Compétence modifiée avec succès.")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de la modification : {str(e)}")

    

    def on_emploi_selectionne(self, item, column):
        # Vérifie que c’est bien un item de niveau 2 (emploi)
        parent = item.parent()
        if parent and parent.parent():  # donc niveau 2
            emploi_nom = item.text(0)
            emploi = self.job_controller.get_emploi_par_nom(emploi_nom)
            self.selected_emploi = emploi["id"] if emploi else None

        else:
            self.selected_emploi = None
    def sauvegarder_competences_emploi_table(self):
        if not hasattr(self, 'selected_emploi') or not self.selected_emploi:
            QMessageBox.warning(self, "Aucun emploi", "Veuillez sélectionner un emploi.")
            return

        table = self.ui.table_competences
        competences = []
        print('yes')
        for row in range(table.rowCount()):
            nom = table.item(row, 0).text().strip()
            print('yes 1')
            niveau = table.item(row, 1).text().strip()
            print('yes 2')
            type_= table.item(row, 2).text().strip()
            print('yes 3')

            competences.append({
                "nom": nom,
                "type": type_,
                "niveau_requis": niveau
            })

        try:
            print('competences avant controller', competences)
            self.competence_controller.sauvegarder_competences_emploi(self.selected_emploi, competences)
            QMessageBox.information(self, "Succès", "Compétences sauvegardées avec succès.")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de la sauvegarde : {str(e)}")

    def exporter_carto_competences_emploi(self):
        filepath, _ = QFileDialog.getSaveFileName(self, "Exporter Cartographie Compétences", "", "Excel (*.xlsx)")
        if not filepath:
            return

        try:
            import xlsxwriter

            # Récupérer tous les emplois
            emplois = self.job_controller.get_liste_emplois()

            workbook = xlsxwriter.Workbook(filepath)
            worksheet = workbook.add_worksheet("Carto Compétences")

            # En-têtes
            worksheet.write(0, 0, "Emploi")
            worksheet.write(0, 1, "Compétence")

            row = 1
            for emploi in emplois:
                emploi_nom = emploi["titre"]
                competences = crud.get_competences_by_emploi(emploi["id"])
                if not competences:
                    worksheet.write(row, 0, emploi_nom)
                    worksheet.write(row, 1, "(Aucune compétence)")
                    row += 1
                else:
                    for comp in competences:
                        worksheet.write(row, 0, emploi_nom)
                        worksheet.write(row, 1, comp[0])  # comp[0] = nom de la compétence
                        row += 1

            workbook.close()
            QMessageBox.information(self, "Succès", f"Export réalisé :\n{filepath}")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur export : {str(e)}")

    def charger_competences_dans_table(self, emploi_id):
        self.ui.table_competences.setRowCount(0)
        competences = crud.get_competences_by_emploi(emploi_id)

        for row_idx, (nom, niveau, type_comp) in enumerate(competences):
            self.ui.table_competences.insertRow(row_idx)
            self.ui.table_competences.setItem(row_idx, 0, QTableWidgetItem(nom))
            self.ui.table_competences.setItem(row_idx, 1, QTableWidgetItem(str(niveau) if niveau else ""))
            self.ui.table_competences.setItem(row_idx, 2, QTableWidgetItem(type_comp or ""))
    def mettre_a_jour_autocompletion_competence(self, texte):
        if len(texte) < 3 or not self.current_projet_id:
            return

        noms = self.competence_controller.get_autocomplete_competences(
            self.current_projet_id, texte
        )

        print(f"[DEBUG] Compétences récupérées : {noms}")
        self.model_competences.setStringList(noms)

    def ajouter_competence_a_emploi(self):
        nom = self.ui.input_competence.text().strip()
        niveau = self.ui.combo_niveau_requis.currentText().strip()
        type_comp = self.ui.combo_type_competence.currentText().strip()
        definition = self.ui.input_definition_niveau_competence.toPlainText().strip()

        if not nom or not niveau:
            QMessageBox.warning(self, "Champs manquants", "Merci de renseigner la compétence et le niveau.")
            return

        if not hasattr(self, 'selected_emploi') or not self.selected_emploi:
            QMessageBox.warning(self, "Aucun emploi sélectionné", "Veuillez sélectionner un emploi.")
            return

        emploi_id = self.selected_emploi

        # 🔁 Appel via controller
        self.competence_controller.ajouter_competence(
            self.current_projet_id, emploi_id, nom, type_comp, niveau, definition
        )

        # UI update
        row_idx = self.ui.table_competences.rowCount()
        self.ui.table_competences.insertRow(row_idx)
        self.ui.table_competences.setItem(row_idx, 0, QTableWidgetItem(nom))
        self.ui.table_competences.setItem(row_idx, 1, QTableWidgetItem(niveau))
        self.ui.table_competences.setItem(row_idx, 2, QTableWidgetItem(type_comp))

        self.ui.input_competence.clear()
        self.ui.input_definition_niveau_competence.clear()

    def sauvegarder_fiche_de_poste(self):
        # Vérifie si un emploi a été sélectionné auparavant
        if not hasattr(self, 'selected_emploi') or not self.selected_emploi:
            QMessageBox.warning(self, "Sélection requise", "Veuillez sélectionner un emploi.")
            return

        emploi_id = self.selected_emploi

        
        finalite = self.ui.input_finalite_poste.toPlainText().strip()
        
        if not finalite:
            QMessageBox.warning(self, "Finalité manquante", "Veuillez saisir ou générer une finalité avant de sauvegarder.")
            return

        # ⬇️ 1. Sauvegarder la finalité via le controller
        if self.finalite_controller:  # Vérifie que le controller est bien initialisé
            self.finalite_controller.sauvegarder_finalite(emploi_id, finalite)
        else:
            QMessageBox.critical(self, "Erreur", "Le contrôleur de finalité n'est pas initialisé.")

        #self.finalite_controller.sauvegarder_finalite(emploi_id, finalite)

        # ⬇️ 2. Sauvegarder les activités ##### AVEV
        macro_items = [self.ui.tree_activites_fichedeposte.topLevelItem(i) for i in range(self.ui.tree_activites_fichedeposte.topLevelItemCount())]

        activites = []
        for macro_item in macro_items:
            macro_titre = macro_item.text(0).replace("🟦", "").strip()
            micro_activites = []
            for j in range(macro_item.childCount()):
                micro_texte = macro_item.child(j).text(0).replace("🔹", "").strip()
                micro_activites.append(micro_texte)
            activites.append({
                "titre": macro_titre,
                "micro_activites": micro_activites
            })

        # Appelle un controller/ service dédié aux fiches de poste (à créer)
        self.fiche_poste_controller.sauvegarder_activites(emploi_id, activites)

        QMessageBox.information(self, "Succès", "Fiche de poste sauvegardée.")
    def export_fiches(self):
        export_fiches(self) 
    def generer_fiche_de_poste_via_ia(self):
        if not self.current_projet_id:
            QMessageBox.warning(self, "Projet requis", "Veuillez ouvrir un projet.")
            return

        selected_item = self.ui.tree_carto_emploi_fichedeposte.currentItem()
        if not selected_item or not selected_item.parent() or not selected_item.parent().parent():
            QMessageBox.warning(self, "Sélection requise", "Veuillez sélectionner un emploi.")
            return

        emploi_nom = selected_item.text(0)
        emploi = self.job_controller.get_emploi_par_nom(emploi_nom)
        if not emploi:
            QMessageBox.critical(self, "Erreur", f"Emploi introuvable : {emploi_nom}")
            return
  
        emploi_id = emploi["id"]
        # Garde la référence de l'emploi sélectionné
        self.selected_emploi = emploi_id
        projet = crud.get_projet_by_id(self.current_projet_id)
        competences = crud.get_competences_by_emploi(emploi_id)

        secteur = projet["secteur_activite"]
        type_entreprise = projet["type_entreprise"]
        taille = projet["taille"]
        filiales = projet["nb_filiales"]
        international = projet["est_international"]
        digital = projet["niveau_digitalisation"]

        prompt = f"""
    Tu es un expert RH.

    🎯 Objectif : Génère une fiche de poste complète pour l'emploi suivant :

    🧾 Titre du poste : {emploi['titre']}

    📌 Contexte entreprise :
    - Secteur : {secteur}
    - Type d’entreprise : {type_entreprise}
    - Taille : {taille}
    - Nombre de filiales : {filiales}
    - Portée : {"internationale" if international else "nationale"}
    - Niveau de digitalisation : {digital}/10

    🧠 Compétences associées à ce poste :
    {json.dumps([c["nom"] for c in competences], ensure_ascii=False)}

    ---

    🎁 Format attendu (JSON uniquement) :
    {{
        "finalite": "Phrase ou paragraphe résumant la finalité du poste",
        "macro_activites": [
            {{
                "titre": "Nom de la macro-activité",
                "micro_activites": [
                    "Micro-activité 1",
                    "Micro-activité 2",
                    ...
                ]
            }},
            ...
        ]
    }}

    ⛔ Ne donne que du JSON brut. Pas de texte explicatif autour.
    """

        print("[📤 Prompt envoyé à l'IA]:\n", prompt)
        try:
            reponse = self.openai_controller.openai_service.ask_once(prompt)
            texte_json = self.extraire_json_depuis_texte(reponse)
            data = json.loads(texte_json)

            finalite = data.get("finalite", "")
            macros = data.get("macro_activites", [])

            # 👉 Affiche la finalité dans le QTextEdit
            self.ui.input_finalite_poste.clear()
            self.ui.input_finalite_poste.setPlainText(finalite)

            self.ui.tree_activites_fichedeposte.clear()

            for macro in macros:
                titre_macro = macro.get("titre", "").strip()
                micro_activites = macro.get("micro_activites", [])

                if not titre_macro:
                    continue

                macro_item = QtWidgets.QTreeWidgetItem([f"🟦 {titre_macro}"])
                for micro in micro_activites:
                    micro_item = QtWidgets.QTreeWidgetItem([f"🔹 {micro.strip()}"])
                    macro_item.addChild(micro_item)

                self.ui.tree_activites_fichedeposte.addTopLevelItem(macro_item)

            self.ui.tree_activites_fichedeposte.expandAll()
            #crud.set_finalite_emploi(emploi_id, finalite)  # méthode à créer si tu stockes la finalité

            # Actualiser le tree
            #self.charger_fiche_de_poste_tree(emploi_id)

            QMessageBox.information(self, "Succès", "Fiche de poste générée avec succès.")
        
        except Exception as e:
            print("[Erreur IA fiche de poste]", e)
            QMessageBox.critical(self, "Erreur", f"Erreur lors de la génération IA :\n{str(e)}")
    # Fonction pour récupérer les emplois remplis
    
    def importer_cartographie_depuis_excel(self):
        print('ok')
        
        filepath, _ = QFileDialog.getOpenFileName(self, "Importer une cartographie", "", "Excel Files (*.xlsx *.xls)")
        if not filepath:
            print('okkk')
            return

        try:
            df = pd.read_excel(filepath)

            # Vérification des colonnes attendues
            expected_cols = ["Famille", "Sous-famille", "Emploi"]
            if not all(col in df.columns for col in expected_cols):
                QMessageBox.warning(self, "Format incorrect", f"Le fichier doit contenir les colonnes : {', '.join(expected_cols)}")
                return

            # Construction de la structure hiérarchique
            data = {"familles": []}
            familles_dict = {}

            for _, row in df.iterrows():
                fam = str(row["Famille"]).strip()
                sf = str(row["Sous-famille"]).strip()
                emp = str(row["Emploi"]).strip()

                if fam not in familles_dict:
                    familles_dict[fam] = {"nom": fam, "sous_familles": {}}
                if sf and sf not in familles_dict[fam]["sous_familles"]:
                    familles_dict[fam]["sous_familles"][sf] = []

                if sf and emp:
                    familles_dict[fam]["sous_familles"][sf].append(emp)

            for fam_data in familles_dict.values():
                sous_familles = []
                for sf_nom, emplois in fam_data["sous_familles"].items():
                    sous_familles.append({
                        "nom": sf_nom,
                        "emplois": emplois
                    })
                data["familles"].append({
                    "nom": fam_data["nom"],
                    "sous_familles": sous_familles
                })

            # Affichage dans le tree
            self.afficher_hierarchie_dans_tree(data)
            QMessageBox.information(self, "Import terminé", "Cartographie importée avec succès.")

        except Exception as e:
            print("[Erreur import Excel] :", e)
            QMessageBox.critical(self, "Erreur", f"Erreur lors de l’import : {str(e)}")



   
        
    def charger_finalite_poste(self, emploi_id):
        finalite = crud.get_finalite_by_emploi(emploi_id) or ""
        self.ui.input_finalite_poste.clear()
        self.ui.input_finalite_poste.setPlainText(finalite)

    def on_selection_changed(self):
        selected_items = self.ui.tree_carto_emploi_fichedeposte.selectedItems()

        if not selected_items:
            self.ui.btn_lancer_generateur_activites.setEnabled(False)
            return

        selected_item = selected_items[0]

        # Vérifie si l’item est un emploi (niveau 3)
        # Par exemple, si l’emploi est toujours à profondeur 2 (racine = 0)
        if selected_item and selected_item.parent() and selected_item.parent().parent():
            self.ui.btn_lancer_generateur_activites.setEnabled(True)
        else:
            self.ui.btn_lancer_generateur_activites.setEnabled(False)
            
    def charger_competences_dans_fiche_poste(self, emploi_id):
        self.ui.table_competences_fichedeposte.setRowCount(0)
        competences = crud.get_competences_by_emploi(emploi_id)

        for row_idx, (nom, niveau, type_comp) in enumerate(competences):
            self.ui.table_competences_fichedeposte.insertRow(row_idx)
            self.ui.table_competences_fichedeposte.setItem(row_idx, 0, QTableWidgetItem(nom))
            self.ui.table_competences_fichedeposte.setItem(row_idx, 1, QTableWidgetItem(str(niveau) if niveau else ""))
            self.ui.table_competences_fichedeposte.setItem(row_idx, 2, QTableWidgetItem(type_comp or ""))

    def charger_competences_sous_famille(self):
        nom_sous_famille = self.ui.combo_sous_famille.currentText()
        if not nom_sous_famille:
            return

        donnees = self.job_controller.get_competence_table_for_sous_famille(nom_sous_famille)

        # Préparer colonnes (emplois) et lignes (compétences)
        emplois = sorted(set(row["emploi_titre"] for row in donnees if row["emploi_titre"]))
        competences = sorted(set(row["competence_nom"] for row in donnees if row["competence_nom"]))

        # Création d'un dictionnaire (emploi -> compétence -> niveau)
        matrice = {
            comp: {emp: "" for emp in emplois}
            for comp in competences
        }

        for row in donnees:
            emp = row["emploi_titre"]
            comp = row["competence_nom"]
            niveau = row["niveau"]
            if emp and comp:
                matrice[comp][emp] = niveau or ""

        # Mise à jour de la QTableWidget
        table = self.ui.table_niveaux_competence
        table.clear()
        table.setRowCount(len(competences))
        table.setColumnCount(len(emplois))
        table.setHorizontalHeaderLabels(emplois)
        table.setVerticalHeaderLabels(competences)

        for row_idx, comp in enumerate(competences):
            for col_idx, emp in enumerate(emplois):
                niveau = matrice[comp][emp]
                item = QTableWidgetItem(niveau)
                table.setItem(row_idx, col_idx, item)

    def update_combo_sous_famille(self):
        self.ui.combo_sous_famille.clear()

        if self.current_projet_id:
            sous_familles = self.job_controller.get_all_sous_familles()
            self.ui.combo_sous_famille.addItems(sous_familles)
    def on_tree_generateur_item_clicked(self, item):
        # On vérifie que c’est bien un emploi (donc un item de niveau 2)
        parent = item.parent()
        if parent and parent.parent():  # niveau 2 = emploi
            emploi_nom = item.text(0)
            emploi = self.job_controller.get_emploi_par_nom(emploi_nom)
            if emploi:
                self.charger_competences_dans_table(emploi["id"])

    def on_tree_fiche_poste_item_clicked(self, item):
        # On vérifie que c’est bien un emploi (donc un item de niveau 2)
        parent = item.parent()
        if parent and parent.parent():  # niveau 2 = emploi
            emploi_nom = item.text(0)
            emploi = self.job_controller.get_emploi_par_nom(emploi_nom)
            if emploi:
                self.charger_competences_dans_fiche_poste(emploi["id"])

    def charger_dictionnaire_depuis_base(self):
        if not self.current_projet_id:
            return

        # Charger les niveaux
        niveaux = crud.get_niveaux_by_projet(self.current_projet_id)
        print('niveaux',niveaux)
        # Charger les compétences
        competences = crud.get_all_competences_by_projet(self.current_projet_id)

        self.ui.table_dictionnaire_competences.clear()

        headers = ["Compétence", "Définition"] + [f"Niveau {n['position']}" for n in niveaux]
        self.ui.table_dictionnaire_competences.setColumnCount(len(headers))
        self.ui.table_dictionnaire_competences.setHorizontalHeaderLabels(headers)
        self.ui.table_dictionnaire_competences.setRowCount(0)

        for comp in competences:
            row = self.ui.table_dictionnaire_competences.rowCount()
            self.ui.table_dictionnaire_competences.insertRow(row)

            self.ui.table_dictionnaire_competences.setItem(row, 0, QTableWidgetItem(comp["nom"]))
            self.ui.table_dictionnaire_competences.setItem(row, 1, QTableWidgetItem(comp.get("dictionnaire") or ""))

            # Récupérer les descriptions par niveau
            descs = crud.get_descriptions_by_competence(comp["id"])
            niveau_map = {d["niveau_id"]: d["description"] for d in descs}

            for i, niveau in enumerate(niveaux):
                desc = niveau_map.get(niveau["id"], "")
                self.ui.table_dictionnaire_competences.setItem(row, 2 + i, QTableWidgetItem(desc))
    
    def ajouter_ligne_dictionnaire(self):
        row = self.ui.table_dictionnaire_competences.rowCount()
        self.ui.table_dictionnaire_competences.insertRow(row)

    def supprimer_ligne_dictionnaire(self):
        rows = sorted(set(index.row() for index in self.ui.table_dictionnaire_competences.selectedIndexes()), reverse=True)
        for r in rows:
            self.ui.table_dictionnaire_competences.removeRow(r)
    def generer_dictionnaire_via_ia(self):
        if not self.current_projet_id:
            QMessageBox.warning(self, "Projet requis", "Veuillez ouvrir un projet.")
            return

        projet = crud.get_projet_by_id(self.current_projet_id)
        carto = self.extraire_cartographie_depuis_tree()
        referentiel = crud.get_referentiel_competences_by_projet(self.current_projet_id)
        niveaux = crud.get_niveaux_by_projet(self.current_projet_id)
        dict_def_niveaux = {str(n["position"]): n["definition"] for n in niveaux if n["definition"]}

        nb_niveaux = projet["nb_niveaux"] or len(niveaux)
        secteur = projet["secteur_activite"]
        type_entreprise = projet["type_entreprise"]
        taille = projet["taille"]
        filiales = projet["nb_filiales"]
        international = projet["est_international"]
        digital = projet["niveau_digitalisation"]
        nb_comp = self.ui.spin_nb_competences_projet_dictionnaire.value() if hasattr(self.ui, "spin_nb_competences_dictionnaire") else 10
        # 👉 Construction du prompt
        prompt = f"""
    Tu es un expert RH.

    Génère un dictionnaire de compétences pour une entreprise avec le contexte suivant :

    📌 Contexte :
    - Secteur : {secteur}
    - Type d’entreprise : {type_entreprise}
    - Taille : {taille}
    - Nombre de filiales : {filiales}
    - Portée : {"internationale" if international else "nationale"}
    - Niveau de digitalisation : {digital}/10

    🗺️ Cartographie des emplois :
    {json.dumps(carto, ensure_ascii=False) if carto else "Aucune cartographie fournie"}

    📚 Référentiel existant :
    {json.dumps(referentiel, ensure_ascii=False) if referentiel else "Aucun"}

    🧭 Si aucun métier n'est fourni, imagine une structure RH typique pour ce type d’entreprise : métiers support, métiers techniques, commerciaux, logistique, etc.

    📌 Types de compétences attendues :
    - Compétences cœur métier (ex : audit, comptabilité, gestion de crise)
    - Compétences managériales (leadership, stratégie)
    - Compétences transverses (communication, collaboration, numérique)
    
    🎯 Nombre de niveaux : {nb_niveaux}
    📖 Définitions générales des niveaux : 
    {json.dumps(dict_def_niveaux, ensure_ascii=False) if dict_def_niveaux else "Non fourni"}

    ---

    🔁 Format attendu :
    [
    {{
        "nom": "Compétence X",
        "definition": "Définition globale",
        "niveaux": {{
        "1": "Définition niveau 1",
        "2": "Définition niveau 2",
        ...
        }}
    }},
    ...
    ]

    ⛔ Ne donne que du JSON brut.
    """
        prompt += f"\n📌 Génère environ {nb_comp} compétences pertinentes selon les critères ci-dessus."
        print("[📤 Prompt envoyé à l'IA]:\n", prompt)
        try:
            reponse = self.openai_controller.openai_service.ask_once(prompt)
            texte_json = self.extraire_json_depuis_texte(reponse)
            data = json.loads(texte_json)

            # 🧠 Sauvegarde dans la base
            position_to_id = {str(n["position"]): n["id"] for n in niveaux}
            #crud.supprimer_competences_referentiel(self.current_projet_id)

            for ligne in data:
                nom = ligne.get("nom", "").strip()
                definition = ligne.get("definition", "").strip()
                niveaux_dict = ligne.get("niveaux", {})

                if not nom:
                    continue

                crud.add_competence(nom, self.current_projet_id, description=definition)
                comp = crud.get_competence_by_nom(nom, self.current_projet_id)
                if not comp:
                    continue

                for position, texte in niveaux_dict.items():
                    niveau_id = position_to_id.get(str(position))
                    if niveau_id and texte:
                        crud.add_competence_niveau_description(comp["id"], niveau_id, texte.strip())

            # 🔄 Afficher dans la table
            self.charger_dictionnaire_depuis_base()

            # ✅ Mettre à jour liste du référentiel aussi (synchro)
            liste_referentiel = crud.get_referentiel_competences_by_projet(self.current_projet_id)
            self.afficher_referentiel_dans_liste(liste_referentiel)

            QMessageBox.information(self, "Succès", "Dictionnaire généré et enregistré avec succès.")

        except Exception as e:
            print("[Erreur IA dictionnaire]", e)
            QMessageBox.critical(self, "Erreur", f"Erreur lors de la génération IA :\n{str(e)}")

    def charger_definition_niveau_selectionne(self):
        niveau_str = self.ui.combo_niveau_competence_a_definir.currentText()

        if not niveau_str:
            #print("📛 Niveau non sélectionné")
            self.ui.input_definition_niveau_competence_2.clear()
            return

        try:
            niveau_int = int(niveau_str)
        except ValueError:
            #print("📛 Niveau invalide")
            self.ui.input_definition_niveau_competence_2.clear()
            return

        niveaux = crud.get_niveaux_by_projet(self.current_projet_id)
        #print("🔎 Niveaux projet :", niveaux)
        niveau_obj = next((n for n in niveaux if n["position"] == niveau_int), None)

        if not niveau_obj:
            #print(f"📛 Aucun niveau avec position {niveau_int}")
            self.ui.input_definition_niveau_competence_2.clear()
            return

        texte = crud.get_definition_niveau(niveau_obj["id"]) or ""
        #print(f"📥 Texte récupéré pour niveau {niveau_int} (id {niveau_obj['id']}) : {texte}")
        self.ui.input_definition_niveau_competence_2.setPlainText(texte)



    def sauvegarder_definition_niveau(self):
        niveau_str = self.ui.combo_niveau_competence_a_definir.currentText().strip()
        texte = self.ui.input_definition_niveau_competence_2.toPlainText().strip()

        if not niveau_str:
            QMessageBox.warning(self, "Erreur", "Veuillez sélectionner un niveau.")
            return

        try:
            niveau_int = int(niveau_str)
        except ValueError:
            QMessageBox.warning(self, "Erreur", "Sélection invalide.")
            return

        niveaux = crud.get_niveaux_by_projet(self.current_projet_id)
        niveau_obj = next((n for n in niveaux if n["position"] == niveau_int), None)

        if not niveau_obj:
            QMessageBox.warning(self, "Erreur", "Niveau introuvable.")
            return

        #print(f"📝 Sauvegarde : niveau {niveau_obj['id']} -> {texte}")
        crud.update_niveau_definition(niveau_obj["id"], texte)
        QMessageBox.information(self, "Succès", "Définition sauvegardée.")

    
    def sauvegarder_dictionnaire(self):
        table = self.ui.table_dictionnaire_competences
        niveaux = crud.get_niveaux_by_projet(self.current_projet_id)

        crud.supprimer_competences_referentiel(self.current_projet_id)

        for row in range(table.rowCount()):
            nom = table.item(row, 0).text().strip()
            definition = table.item(row, 1).text().strip()
            crud.add_competence(nom, self.current_projet_id, description=definition)

            comp = crud.get_competence_by_nom(nom, self.current_projet_id)
            if not comp:
                continue
            comp_id = comp["id"]

            for i, niveau in enumerate(niveaux):
                col = 2 + i
                if col < table.columnCount():
                    desc = table.item(row, col)
                    if desc:
                        crud.add_competence_niveau_description(comp_id, niveau["id"], desc.text().strip())

    def importer_dictionnaire_excel(self):
        filepath, _ = QFileDialog.getOpenFileName(self, "Importer un dictionnaire", "", "Excel Files (*.xlsx *.xls)")
        if not filepath:
            return

        try:
            df = pd.read_excel(filepath)
            if df.empty:
                QMessageBox.warning(self, "Fichier vide", "Le fichier ne contient aucune donnée.")
                return

            self.ui.table_dictionnaire_competences.clear()
            self.ui.table_dictionnaire_competences.setColumnCount(len(df.columns))
            self.ui.table_dictionnaire_competences.setHorizontalHeaderLabels(df.columns)

            self.ui.table_dictionnaire_competences.setRowCount(0)
            for row_idx, row in df.iterrows():
                self.ui.table_dictionnaire_competences.insertRow(row_idx)
                for col_idx, value in enumerate(row):
                    item = QTableWidgetItem(str(value) if pd.notna(value) else "")
                    item.setTextAlignment(QtCore.Qt.AlignTop | QtCore.Qt.AlignLeft)
                    item.setFlags(item.flags() ^ QtCore.Qt.ItemIsEditable)  # facultatif si tu veux bloquer l'édition
                    item.setToolTip(item.text())  # pour que l'utilisateur voie tout
                    item.setData(QtCore.Qt.DisplayRole, texte)
                    item.setData(QtCore.Qt.EditRole, texte)
                    item.setTextAlignment(QtCore.Qt.AlignTop)
                    item.setWhatsThis(texte)
                    item.setFont(self.ui.font())  # optionnel
                    item.setText(texte)
                    item.setWrapMode(True)  # ⛔️ ← à ne PAS utiliser ! ça n'existe pas pour QTableWidgetItem

                    # Utilise plutôt :
                    item.setData(QtCore.Qt.DisplayRole, texte)
                    self.ui.table_dictionnaire_competences.setItem(row_idx, col_idx, item)
            self.ui.table_dictionnaire_competences.resizeRowsToContents()

            QMessageBox.information(self, "Import réussi", f"{len(df)} lignes importées depuis le dictionnaire Excel.")

        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de l'import : {str(e)}")
    
    def exporter_dictionnaire_excel(self):
        filepath, _ = QFileDialog.getSaveFileName(self, "Exporter le dictionnaire", "", "Excel Files (*.xlsx)")
        if not filepath:
            return

        try:
            table = self.ui.table_dictionnaire_competences
            row_count = table.rowCount()
            col_count = table.columnCount()

            data = []
            headers = [table.horizontalHeaderItem(i).text() for i in range(col_count)]

            for row in range(row_count):
                row_data = {}
                for col in range(col_count):
                    item = table.item(row, col)
                    row_data[headers[col]] = item.text() if item else ""
                data.append(row_data)

            df = pd.DataFrame(data)
            df.to_excel(filepath, index=False)

            QMessageBox.information(self, "Export terminé", f"Dictionnaire exporté vers :\n{filepath}")

        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de l’export : {str(e)}")

    def exporter_cartographie_en_excel(self):
        if not self.current_projet_id:
            QMessageBox.warning(self, "Projet requis", "Veuillez ouvrir un projet.")
            return

        # Choix du fichier de sortie
        filepath, _ = QFileDialog.getSaveFileName(self, "Exporter la cartographie", "", "Excel Files (*.xlsx)")
        if not filepath:
            return

        try:
            # Préparation des données depuis le tree
            lignes = []
            tree = self.ui.tree_cartographie_metiers

            for i in range(tree.topLevelItemCount()):
                item_famille = tree.topLevelItem(i)
                nom_famille = item_famille.text(0)

                if item_famille.childCount() == 0:
                    lignes.append([nom_famille, "", ""])
                else:
                    for j in range(item_famille.childCount()):
                        item_sf = item_famille.child(j)
                        nom_sf = item_sf.text(0)

                        if item_sf.childCount() == 0:
                            lignes.append([nom_famille, nom_sf, ""])
                        else:
                            for k in range(item_sf.childCount()):
                                item_emploi = item_sf.child(k)
                                nom_emploi = item_emploi.text(0)
                                lignes.append([nom_famille, nom_sf, nom_emploi])

            # Création du fichier Excel
            workbook = xlsxwriter.Workbook(filepath)
            worksheet = workbook.add_worksheet("Cartographie")

            # En-têtes
            headers = ["Famille", "Sous-famille", "Emploi"]
            for col, h in enumerate(headers):
                worksheet.write(0, col, h)

            # Données
            for row, ligne in enumerate(lignes, start=1):
                for col, valeur in enumerate(ligne):
                    worksheet.write(row, col, valeur)

            workbook.close()
            QMessageBox.information(self, "Export réussi", f"La cartographie a été exportée vers :\n{filepath}")

        except Exception as e:
            print("[Erreur export Excel] :", e)
            QMessageBox.critical(self, "Erreur", f"Erreur lors de l’export : {str(e)}")


    def mettre_a_jour_etat_boutons_tree(self):
        item = self.ui.tree_cartographie_metiers.currentItem()
        
        # Par défaut, on désactive tout
        self.ui.btn_ajouter_sous_familles.setEnabled(False)
        self.ui.btn_ajouter_metier.setEnabled(False)

        if item:
            niveau = 0
            parent = item.parent()
            while parent:
                niveau += 1
                parent = parent.parent()

            if niveau == 0:
                # Famille sélectionnée
                self.ui.btn_ajouter_sous_familles.setEnabled(True)
            elif niveau == 1:
                # Sous-famille sélectionnée
                self.ui.btn_ajouter_metier.setEnabled(True)

    def extraire_json_depuis_texte(self,texte: str) -> str:
        """Extrait la première portion JSON valide depuis une réponse contenant du Markdown."""
        try:
            matches = re.findall(r"```json\n(.*?)\n```", texte, re.DOTALL)
            if matches:
                return matches[0].strip()
            else:
                return texte.strip()  # fallback si pas de bloc markdown
        except Exception as e:
            print("[Erreur extraction JSON] :", e)
            return texte

    # --- Methode pour choisir dossier ---
    def choisir_dossier_fichiers(self):
        dossier = QFileDialog.getExistingDirectory(self, "Choisir un dossier")
        if dossier:
            self.chemin_dossier_fichiers = dossier
            QMessageBox.information(self, "Dossier sélectionné", f"{dossier}")

    
    def lancer_generation_competences(self):
        from ui.dialogs.dialog_generation_competences_emploi import DialogGenerationCompetencesParEmploi
        from os import listdir
        from os.path import isfile, join
        import json
        self.interruption_generation = False
        self.ui.btn_ajouter_zone_fichier.setEnabled(True)

        print("🚀 Lancement de la génération de compétences...")

        emplois_existants = self.job_controller.get_liste_emplois()
        emplois_traite_ids = self.competence_controller.get_emploi_ids_avec_competences()
        print('emplois_traite_ids',emplois_traite_ids)
        fichiers_traites = crud.get_fichiers_deja_traite(self.current_projet_id)
        emplois_restants = [emp for emp in emplois_existants if emp["id"] not in emplois_traite_ids]

        # Cas 1 : Fichier unique à structure libre (extraction IA des emplois et compétences)
        if self.chemin_fichier_unique:
            print(f"📄 Fichier unique détecté : {self.chemin_fichier_unique}")
            QMessageBox.information(self, "Traitement fichier", f"Nous commençons à traiter les données extraites du fichier : {self.chemin_fichier_unique}.")

            contenu = self.lire_contenu_fichier(self.chemin_fichier_unique)
            print('contenu',contenu)
            description = self.description_fichier_unique
            prompt_extraction = self.construire_prompt_extraction_depuis_fichier_unique(contenu)
            try:
                print('prompt_extraction',prompt_extraction)
                reponse = self.openai_controller.openai_service.ask_once(prompt_extraction)
                print('response fichier extraite', reponse)
                extractions = json.loads(self.extraire_json_depuis_texte(reponse))
                print('extractions', extractions)
                
            except Exception as e:
                QMessageBox.critical(self, "Erreur IA", f"Erreur extraction depuis fichier : {str(e)}")
                extractions = []

            for extrait in extractions:
                print('entree boucle')
                print('extrait',extrait)
                print('extrait["emploi"]', extrait["emploi"])
                print('extrait["competences"]', extrait["competences"])
                nom_emploi = extrait["emploi"]
                competences_detectees = extrait["competences"]
                self.hilighter_emploi_dans_tree(nom_emploi)
                prompt_matching = self.construire_prompt_matching_emploi(nom_emploi, competences_detectees)
                reponse_matching = self.openai_controller.openai_service.ask_once(prompt_matching)
                print('@@suggestion ia: ', reponse_matching)
                suggestions_ia = json.loads(self.extraire_json_depuis_texte(reponse_matching))
                source =f"Données extraites du fichier : {self.chemin_fichier_unique}."

                dlg = DialogGenerationCompetencesParEmploi(
                    titre_fichier=nom_emploi,
                    contenu_fichier=None,
                    suggestions_ia=suggestions_ia,
                    emplois_existants=emplois_existants,
                    openai_controller=self.openai_controller,
                    prompt_builder=self.construire_prompt_competences,
                    json_extractor=self.extraire_json_depuis_texte,
                    competences_detectees=competences_detectees,
                    source = source,
                    parent = self
                )
                if self.interruption_generation:
                    print("[🛑 Traitement interrompu par l'utilisateur]")
                    return
                print(f"📤 Ouverture du dialog pour : {nom_emploi} (fichier unique)")
                if dlg.exec_() == QDialog.Accepted:
                    donnees = dlg.get_donnees_valides()
                    emploi_nom = dlg.get_nom_emploi()
                    if dlg.combo_emploi.currentText() == "Créer un nouvel emploi...":
                        titre = dlg.input_titre.text().strip()
                        famille = dlg.input_famille.currentText().strip()
                        sous_famille = dlg.input_sous_famille.currentText().strip()
                        if titre and famille and sous_famille:
                            self.job_controller.ajouter_emploi_si_nouveau(titre, famille, sous_famille)
                    # 🔄 Recharge des 3 trees cartographiques
                    if self.job_controller:
                        data = self.job_controller.recuperer_hierarchie_as_json()
                        if data:
                            self.afficher_hierarchie_dans_tree(data)
                            self.afficher_hierarchie_dans_tree_generateur(data)
                            self.afficher_hierarchie_dans_tree_fiche_poste(data)
                            print("📦 Données à sauvegarder CAS FICHIER :", donnees)
                            print("📦 emploi_nom :", emploi_nom)
                    self.competence_controller.sauvegarder_competences_emploi(emploi_nom, donnees)
                    self.marquer_emploi_comme_traite(emploi_nom)
                elif dlg.arret_demande:
                    competences = self.competence_controller.charger_referentiel()
                    self.afficher_referentiel_dans_liste(competences)
                    self.charger_dictionnaire_depuis_base()
                    print("[🛑 Utilisateur a demandé l’arrêt]")
                    break  # stoppe la boucle proprement
                    return
                
        if self.interruption_generation:
            competences = self.competence_controller.charger_referentiel()
            self.afficher_referentiel_dans_liste(competences)
            self.charger_dictionnaire_depuis_base()
            print("[🛑 Traitement interrompu par l'utilisateur]")
            return
        # Cas 2 : Dossier de fichiers (traitement fichier par fichier)
        if self.chemin_dossier_fichiers:
            print(f"📁 Dossier de fichiers détecté : {self.chemin_dossier_fichiers}")
            QMessageBox.information(self, "Traitement dossier", f"Nous commençons à traiter les fichiers du dossier : {self.chemin_dossier_fichiers}.")
            fichiers = [f for f in listdir(self.chemin_dossier_fichiers) if isfile(join(self.chemin_dossier_fichiers, f))]

            for f in fichiers:
                if f in fichiers_traites:
                    print(f"[⏭️ Fichier déjà traité] {f}")
                    continue

                contenu = self.lire_contenu_fichier(join(self.chemin_dossier_fichiers, f))
                try:
                    prompt = self.construire_prompt_matching_emploi(f, contenu)
                    if not prompt:
                        suggestions = []
                    else:
                        reponse = self.openai_controller.openai_service.ask_once(prompt)
                        suggestions = json.loads(self.extraire_json_depuis_texte(reponse))
                except Exception as e:
                    print(f"[❌ Erreur IA sur {f}] {e}")
                    suggestions = []
                source = f"Données extraites du dossier : {self.chemin_dossier_fichiers}"
                dlg = DialogGenerationCompetencesParEmploi(
                    titre_fichier=f,
                    contenu_fichier=contenu,
                    suggestions_ia=suggestions,
                    emplois_existants=emplois_existants,
                    openai_controller=self.openai_controller,
                    prompt_builder=self.construire_prompt_competences,
                    json_extractor=self.extraire_json_depuis_texte,
                    source = source,
                    parent=self
                )
                if self.interruption_generation:
                    competences = self.competence_controller.charger_referentiel()
                    self.afficher_referentiel_dans_liste(competences)
                    self.charger_dictionnaire_depuis_base()
                    print("[🛑 Traitement interrompu par l'utilisateur]")
                    return
                print(f"📤 Ouverture du dialog pour : {f} | fichier: oui")
                if dlg.exec_() == QDialog.Accepted:
                    donnees = dlg.get_donnees_valides()
                    emploi_nom = dlg.get_nom_emploi()
                    if dlg.combo_emploi.currentText() == "Créer un nouvel emploi...":
                        titre = dlg.input_titre.text().strip()
                        famille = dlg.input_famille.currentText().strip()
                        sous_famille = dlg.input_sous_famille.currentText().strip()
                        if titre and famille and sous_famille:
                            self.job_controller.ajouter_emploi_si_nouveau(titre, famille, sous_famille)
                    # 🔄 Recharge des 3 trees cartographiques
                    if self.job_controller:
                        data = self.job_controller.recuperer_hierarchie_as_json()
                        if data:
                            self.afficher_hierarchie_dans_tree(data)
                            self.afficher_hierarchie_dans_tree_generateur(data)
                            self.afficher_hierarchie_dans_tree_fiche_poste(data)
                    print("📦 Données à sauvegarder CAS DOSSIER :", donnees)
                    print("📦 emploi_nom :", emploi_nom)
                    self.competence_controller.sauvegarder_competences_emploi(emploi_nom, donnees)
                    emploi_id = crud.get_id_emploi_par_nom(emploi_nom)
                    crud.marquer_fichier_comme_traite(f, emploi_id, self.current_projet_id)
                    

                    self.marquer_emploi_comme_traite(emploi_nom)
                elif dlg.arret_demande:
                    print("[🛑 Utilisateur a demandé l’arrêt]")
                    break  # stoppe la boucle proprement
                    return
        if self.interruption_generation:
            print("[🛑 Traitement interrompu par l'utilisateur]")
            competences = self.competence_controller.charger_referentiel()
            self.afficher_referentiel_dans_liste(competences)
            self.charger_dictionnaire_depuis_base()
            return        
        # Cas 3 : Traitement des emplois restants dans la cartographie
        emplois_restants = [emp for emp in self.job_controller.get_liste_emplois() if emp["id"] not in self.competence_controller.get_emploi_ids_avec_competences()]
        if emplois_restants:
            QMessageBox.information(self, "Traitement cartographie", "Nous commençons le traitement des emplois restants dans la cartographie.")
        for emp in emplois_restants:
            titre = emp["titre"]
            self.hilighter_emploi_dans_tree(titre)
            source = "Données extraites de la cartographie des emplois."
            dlg = DialogGenerationCompetencesParEmploi(
                titre_fichier=titre,
                contenu_fichier=None,
                suggestions_ia=[],
                emplois_existants=emplois_existants,
                openai_controller=self.openai_controller,
                prompt_builder=self.construire_prompt_competences,
                json_extractor=self.extraire_json_depuis_texte,
                source = source,
                parent = self
            )
            print(f"📤 Ouverture du dialog pour : {titre} | fichier: non")
            if dlg.exec_() == QDialog.Accepted:
                donnees = dlg.get_donnees_valides()
                emploi_nom = dlg.get_nom_emploi()
                print("📦 Données à sauvegarder CAS CARTO:", donnees)
                print("📦 emploi_nom :", emploi_nom)
                self.competence_controller.sauvegarder_competences_emploi(emploi_nom, donnees)
                self.marquer_emploi_comme_traite(emploi_nom)
            else:
                competences = self.competence_controller.charger_referentiel()
                self.afficher_referentiel_dans_liste(competences)
                self.charger_dictionnaire_depuis_base()
                print("[🛑 Arrêt utilisateur ou annulation]")
                break
        competences = self.competence_controller.charger_referentiel()
        self.afficher_referentiel_dans_liste(competences)
        self.charger_dictionnaire_depuis_base()
    def lancer_generation_competencesv2(self):
        from ui.dialogs.dialog_generation_competences_emploi import DialogGenerationCompetencesParEmploi
        from os import listdir
        from os.path import isfile, join
        import json


        emplois_traite_ids = self.competence_controller.get_emploi_ids_avec_competences()
        emplois_existants = self.job_controller.get_liste_emplois()
        emplois_a_traiter = []
        #fichiers_traites = crud.get_fichiers_traites()
        fichiers_traites = crud.get_fichiers_deja_traite(self.current_projet_id)

        if self.chemin_dossier_fichiers:
            
            fichiers = [f for f in listdir(self.chemin_dossier_fichiers) if isfile(join(self.chemin_dossier_fichiers, f))]
            

            for f in fichiers:
                if f in fichiers_traites:
                    print(f"[⏭️ Fichier déjà traité] {f}")
                    continue
                contenu = self.lire_contenu_fichier(join(self.chemin_dossier_fichiers, f))
                emplois_a_traiter.append((f, contenu))
        else:
            
            for emp in emplois_existants:
                if emp["id"] not in emplois_traite_ids:
                    emplois_a_traiter.append((emp["titre"], None))

        

        for titre_fichier, contenu_fichier in emplois_a_traiter:
            
            self.hilighter_emploi_dans_tree(titre_fichier)

            # 🧠 Matching IA si on a un fichier
            suggestions = []
            if contenu_fichier:
                try:
                    print(f"🔎 Matching IA pour {titre_fichier}")
                    prompt = self.construire_prompt_matching_emploi(titre_fichier, contenu_fichier)
                    if not prompt:
                        suggestions = []  # 👈 skip proprement si aucun emploi dispo
                    reponse = self.openai_controller.openai_service.ask_once(prompt)
                    suggestions = json.loads(self.extraire_json_depuis_texte(reponse))
                except Exception as e:
                    print(f"[❌ Erreur IA sur {titre_fichier}] {e}")
                    QMessageBox.critical(self, "Erreur IA", f"Erreur IA sur {titre_fichier} : {str(e)}")
                    continue

            dlg = DialogGenerationCompetencesParEmploi(
                titre_fichier=titre_fichier,
                contenu_fichier=contenu_fichier,
                suggestions_ia=suggestions,
                emplois_existants=emplois_existants,
                openai_controller=self.openai_controller,
                prompt_builder=self.construire_prompt_competences,
                json_extractor=self.extraire_json_depuis_texte,
                parent=None
            )


            print(f"📤 Ouverture du dialog pour : {titre_fichier} | fichier: {'oui' if contenu_fichier else 'non'}")

            if not contenu_fichier and not any(emp["titre"] == titre_fichier for emp in emplois_existants):
                print(f"[⚠️ Avertissement] {titre_fichier} ne correspond à aucun emploi connu et ne vient pas d’un fichier. Dialog ignoré.")
                continue

            if dlg.exec_() == QDialog.Accepted:
                donnees_valides = dlg.get_donnees_valides()
                emploi_nom = dlg.get_nom_emploi()
                self.competence_controller.sauvegarder_competences_emploi(emploi_nom, donnees_valides)
                if contenu_fichier:  # on vient d’un fichier
                    emploi_id = crud.get_id_emploi_par_nom(emploi_nom)
                    #crud.marquer_fichier_traite(titre_fichier, emploi_id)
                    crud.marquer_fichier_comme_traite(titre_fichier, emploi_id, self.current_projet_id)


                self.marquer_emploi_comme_traite(emploi_nom)
            else:
                print("[🛑 Arrêt utilisateur ou annulation]")
                break

    def lancer_generation_competences_v4(self):
        print("🚀 Lancement de la génération de compétences...")
        from ui.dialogs.dialog_generation_competences_emploi import DialogGenerationCompetencesParEmploi
        import os
 
        emplois_traite_ids = self.competence_controller.get_emploi_ids_avec_competences()
        emplois_existants = self.job_controller.get_liste_emplois()
        print(f"🧾 Emplois existants : {len(emplois_existants)}")
        print(f"✅ Emplois déjà traités : {emplois_traite_ids}")

        emplois_a_traiter = []

        if self.chemin_dossier_fichiers:
            print(f"📂 Lecture des fichiers depuis : {self.chemin_dossier_fichiers}")
            from os import listdir
            from os.path import isfile, join
            fichiers = [f for f in listdir(self.chemin_dossier_fichiers) if isfile(join(self.chemin_dossier_fichiers, f))]
            print(f"🔍 Fichiers trouvés : {fichiers}")

            for f in fichiers:
                print(f"🤖 IA matching pour le fichier : {f}")
                contenu = self.lire_contenu_fichier(join(self.chemin_dossier_fichiers, f))

                # 🧠 IA Matching
                try:
                    prompt_match = self.construire_prompt_matching_emploi(f, contenu)
                    reponse = self.openai_controller.openai_service.ask_once(prompt_match)
                    suggestions = json.loads(self.extraire_json_depuis_texte(reponse))
                except Exception as e:
                    print("[Erreur matching IA]", e)
                    print(f"[❌ Erreur IA sur fichier {f}] {e}")
                    QMessageBox.critical(self, "Erreur IA", f"Erreur IA sur {f} : {str(e)}")
                    continue

                # 📋 Ajout à la liste à traiter
                emplois_a_traiter.append((f, contenu, suggestions))
        else:
            print("📚 Mode base de données (pas de fichiers)")
            # 📚 Emplois en base non traités
            for emp in emplois_existants:
                if emp["id"] not in emplois_traite_ids:
                    print(f"🟢 À traiter : {emp['titre']}")
                    emplois_a_traiter.append((emp["titre"], None, None))

        for titre_fichier, contenu_fichier, suggestions in emplois_a_traiter:
            self.hilighter_emploi_dans_tree(titre_fichier)

            dlg = DialogGenerationCompetencesParEmploi(
                titre_fichier=titre_fichier,
                contenu_fichier=contenu_fichier,
                suggestions_ia=suggestions,
                emplois_existants=emplois_existants,
                parent=self
            )

            if dlg.exec_() == QDialog.Accepted:
                donnees_valides = dlg.get_donnees_valides()
                emploi_nom = dlg.get_nom_emploi()
                self.competence_controller.sauvegarder_competences_emploi(emploi_nom, donnees_valides)
                self.marquer_emploi_comme_traite(emploi_nom)
            else:
                print("[INFO] Génération interrompue")
                break

    
    def hilighter_emploi_dans_tree(self, nom_emploi):
        items = self.ui.tree_cartographie_metier_generateur_competence.findItems(nom_emploi, Qt.MatchExactly | Qt.MatchRecursive)
        if items:
            item = items[0]
            self.ui.tree_cartographie_metier_generateur_competence.setCurrentItem(item)
            item.setBackground(0, QBrush(QColor("lightblue")))
    
    def marquer_emploi_comme_traite(self, nom_emploi):
        items = self.ui.tree_cartographie_metier_generateur_competence.findItems(nom_emploi, Qt.MatchExactly | Qt.MatchRecursive)
        
        if items:
            item = items[0]
            item.setForeground(0, QBrush(QColor("green")))
            item.setIcon(0, QIcon(":/icon/emploi-selectionne2_resized.ico"))  # à condition d'avoir cette icône


    def construire_prompt_matching_emploi(self, titre_fichier, contenu_fichier):
        """
        Demande à l’IA de trouver les 3 emplois les plus proches parmi les NON TRAITÉS.
        """
        # 🔍 Emplois déjà traités
        emplois_traite_ids = self.competence_controller.get_emploi_ids_avec_competences()
        emplois_existants = self.job_controller.get_liste_emplois()
        emploi_id_par_nom = {emp["titre"]: emp["id"] for emp in emplois_existants}

        # 📦 Cartographie complète (liste directe !)
        cartographie = self.job_controller.get_hierarchie_as_dict()

        # ✅ Vérification format attendu
        if not isinstance(cartographie, list):
            raise ValueError("❌ La cartographie attendue est une liste de familles.")

        # ✂️ On filtre les emplois déjà traités
        for famille in cartographie:
            for sous_famille in famille.get("sous_familles", []):
                sous_famille["emplois"] = [
                    emploi for emploi in sous_famille.get("emplois", [])
                    if isinstance(emploi, dict)
                    and emploi.get("titre")
                    and emploi_id_par_nom.get(emploi["titre"]) not in emplois_traite_ids
                ]
        nb_emplois_restants = sum(
        len(sf["emplois"]) for fam in cartographie for sf in fam.get("sous_familles", [])
        )
        print(f"📊 Emplois non traités proposés à l’IA : {nb_emplois_restants}")
        # 🔒 Si tous les emplois sont traités
        if all(not sf["emplois"] for fam in cartographie for sf in fam.get("sous_familles", [])):
            print("⚠️ Aucun emploi non traité dans la cartographie.")
            return None

        # 🧠 Construction du prompt
        prompt = f"""
    Tu es un assistant RH. Voici un extrait d'un fichier de poste nommé \"{titre_fichier}\".
    Ton objectif est de trouver les 3 emplois les plus proches dans la cartographie ci-dessous.
    Attention, tu dois absolument choisir parmi les emplois listés ci-dessous. Ne propose aucun autre emploi, même s’il semble pertinent.
    Si tu n’en trouves aucun pertinent, retourne simplement une liste vide : []
    Fichier :
    \"\"\"
    {contenu_fichier[:3000]}
    \"\"\"

    Cartographie :
    {json.dumps({"familles": cartographie}, indent=2)}

    Retourne uniquement un JSON avec une liste triée par pertinence, comme ceci :
    [
        {{
            "titre": "Responsable Maintenance",
            "famille": "Production",
            "sous_famille": "Maintenance"
        }},
        ...
    ]
    """.strip()

        return prompt


    def construire_prompt_matching_emploi2(self, titre_fichier, contenu_fichier):
        """
        Demande à l’IA de trouver les 3 emplois les plus proches de la cartographie.
        """
        cartographie = self.job_controller.get_hierarchie_as_dict()
        prompt = f"""
    Tu es un assistant RH. Voici un extrait d'un fichier de poste nommé "{titre_fichier}".
    Ton objectif est de trouver les 3 emplois les plus proches dans la cartographie ci-dessous.

    Fichier :
    \"\"\"
    {contenu_fichier[:1000]}
    \"\"\"

    Cartographie :
    {json.dumps(cartographie, indent=2)}

    Retourne un JSON avec une liste triée par pertinence, comme ceci :
    [
        {{
            "titre": "Responsable Maintenance",
            "famille": "Production",
            "sous_famille": "Maintenance"
        }},
        ...
    ]
        """.strip()
        return prompt
    def ajouter_zone_fichier(self):
        frame = QtWidgets.QFrame()
        frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        layout = QtWidgets.QHBoxLayout(frame)
        layout.setContentsMargins(0, 0, 0, 0)

        # 🔹 Champ de description
        input_desc = QtWidgets.QLineEdit()
        input_desc.setPlaceholderText("Décrivez le contenu du fichier (ex: Offre RH)")
        layout.addWidget(input_desc)

        # 🔹 Bouton d'import
        btn_import = QtWidgets.QPushButton("📁 Importer")
        layout.addWidget(btn_import)

        def importer_fichier():
            from PyQt5.QtWidgets import QFileDialog
            path, _ = QFileDialog.getOpenFileName(
                self, "Choisir un fichier", "", "Fichiers (*.txt *.docx *.xlsx *.xls *.csv *.pdf)"
            )
            if path:
                description = input_desc.text().strip() or "Sans description"
                self.chemin_fichier_unique = path
                self.description_fichier_unique = description
                print(f"📂 Fichier sélectionné : {path} ({description})")
                QMessageBox.information(self, "Fichier chargé", f"{os.path.basename(path)} chargé avec succès.")

        btn_import.clicked.connect(importer_fichier)

        # 🔹 Ajouter ce frame dans la zone dédiée
        self.ui.layout_zone_fichiers.addWidget(frame)
        self.ui.btn_ajouter_zone_fichier.setEnabled(False)

    def clear_layout(self, layout):
        while layout.count():
            child = layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

    def afficher_referentiel_dans_liste(self, liste_competences):
        self.ui.list_referentiel_competences.clear()
        for comp in liste_competences:
            item = QListWidgetItem(comp)
            item.setFlags(item.flags() | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEditable)
            item.setCheckState(QtCore.Qt.Unchecked)
            self.ui.list_referentiel_competences.addItem(item)

    def importer_referentiel_competences_excel(self):
        filepath, _ = QFileDialog.getOpenFileName(self, "Importer un fichier Excel", "", "Excel (*.xlsx *.xls)")
        if not filepath:
            return

        try:
            df = pd.read_excel(filepath)
            competences = df.iloc[:, 0].dropna().astype(str).tolist()
            self.afficher_referentiel_dans_liste(competences)
            QMessageBox.information(self, "Succès", f"{len(competences)} compétences importées.")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de l'import : {str(e)}")

    def generer_referentiel_via_ia(self):
        secteur = self.ui.input_secteur_activite.text().strip()
        type_entreprise = self.ui.combo_type_entreprise.currentText().strip()
        carto = self.extraire_cartographie_depuis_tree()
        carto_str = json.dumps(carto, ensure_ascii=False)

        prompt = (
            f"Génère un référentiel de compétences pour une entreprise dans le secteur : {secteur}, de type {type_entreprise}.\n"
            f"Voici la cartographie des emplois : {carto_str}\n"
            "Retourne un tableau JSON simple : [\"compétence 1\", \"compétence 2\", ...]."
        )

        print("[Prompt génération référentiel] :", prompt)

        try:
            reponse = self.openai_controller.openai_service.ask_once(prompt)
            texte_json = self.extraire_json_depuis_texte(reponse)
            liste = json.loads(texte_json)
            self.afficher_referentiel_dans_liste(liste)
            QMessageBox.information(self, "Référentiel généré", f"{len(liste)} compétences générées.")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", str(e))

    def ajouter_competence_referentiel(self):
        text, ok = QInputDialog.getText(self, "Ajouter compétence", "Nom de la compétence :")
        if ok and text.strip():
            item = QListWidgetItem(text.strip())
            item.setFlags(item.flags() | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEditable)
            item.setCheckState(QtCore.Qt.Unchecked)
            self.ui.list_referentiel_competences.addItem(item)

    def supprimer_competences_selectionnees(self):
        selected_items = [item for item in self.ui.list_referentiel_competences.selectedItems()]
        for item in selected_items:
            self.ui.list_referentiel_competences.takeItem(self.ui.list_referentiel_competences.row(item))
    def supprimer_competences_selectionnees_competencesemploi(self):
        if not hasattr(self, 'selected_emploi') or not self.selected_emploi:
            QMessageBox.warning(self, "Aucun emploi", "Veuillez sélectionner un emploi.")
            return

        table = self.ui.table_competences
        selected_rows = sorted(set(index.row() for index in table.selectedIndexes()), reverse=True)

        if not selected_rows:
            QMessageBox.warning(self, "Aucune sélection", "Veuillez sélectionner une ou plusieurs lignes.")
            return

        for row in selected_rows:
            nom_comp = table.item(row, 0).text().strip()
            # ✅ Suppression via controller (à implémenter dans le controller + service)
            self.competence_controller.supprimer_competence_de_emploi_et_referentiel(
                self.selected_emploi,
                nom_comp
            )
            table.removeRow(row)

        QMessageBox.information(self, "Suppression réussie", "Les compétences sélectionnées ont été supprimées.")

    def fusionner_competences_selectionnees(self):
        items = [self.ui.list_referentiel_competences.item(i) for i in range(self.ui.list_referentiel_competences.count())]
        cochées = [item for item in items if item.checkState() == QtCore.Qt.Checked]

        if len(cochées) < 2:
            QMessageBox.warning(self, "Fusion impossible", "Veuillez cocher au moins 2 compétences à fusionner.")
            return

        noms = [item.text() for item in cochées]
        fusion, ok = QInputDialog.getText(self, "Fusion", f"Fusionner {len(noms)} compétences en :")
        if ok and fusion.strip():
            # Supprimer anciennes
            for item in cochées:
                self.ui.list_referentiel_competences.takeItem(self.ui.list_referentiel_competences.row(item))
            # Ajouter nouvelle
            new_item = QListWidgetItem(fusion.strip())
            new_item.setFlags(new_item.flags() | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEditable)
            new_item.setCheckState(QtCore.Qt.Unchecked)
            self.ui.list_referentiel_competences.addItem(new_item)

    def sauvegarder_referentiel_competence(self):
        if not self.current_projet_id:
            QMessageBox.warning(self, "Projet requis", "Veuillez ouvrir un projet.")
            return

        liste = [self.ui.list_referentiel_competences.item(i).text().strip()
                for i in range(self.ui.list_referentiel_competences.count())]

        self.competence_controller.sauvegarder_referentiel(liste)
        QMessageBox.information(self, "Sauvegarde", "Référentiel sauvegardé avec succès.")

    def exporter_referentiel_competence_excel(self):
        filepath, _ = QFileDialog.getSaveFileName(self, "Exporter Référentiel", "", "Excel (*.xlsx)")
        if not filepath:
            return

        try:
            competences = [self.ui.list_referentiel_competences.item(i).text().strip()
                        for i in range(self.ui.list_referentiel_competences.count())]
            df = pd.DataFrame(competences, columns=["Compétence"])
            df.to_excel(filepath, index=False)
            QMessageBox.information(self, "Export réussi", f"{len(competences)} compétences exportées.")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur export : {str(e)}")

    # --- Lire contenu fichier (simple pour .txt, à adapter si .docx, .xlsx etc) ---
    def construire_prompt_extraction_depuis_fichier_unique(self, contenu):
        """
        Construit un prompt pour demander à l'IA d'extraire une liste d'emplois et leurs compétences depuis un fichier unique.
        """
        prompt = f"""
    Tu es un expert en RH. À partir du contenu ci-dessous d'un fichier contenant plusieurs descriptions de postes ou offres,
    tu dois extraire pour chaque emploi :
    - Le nom du poste
    - Une liste de compétences associées (noms uniquement)

    Fichier :
    \"\"\"
    {contenu}
    \"\"\"

    Retourne un JSON au format :
    [
    {{
        "emploi": "Responsable RH",
        "competences": [
        "Communication",
        "Gestion des talents",
        ...
        ]
    }},
    ...
    ]

    ⚠️ Ne génère rien d’autre que ce JSON.
        """.strip()

        return prompt

    def lire_contenu_fichier(self, path):
        import os
        import pandas as pd

        if not os.path.isfile(path):
            print(f"[ERREUR] Le fichier n'existe pas : {path}")
            return ""

        extension = os.path.splitext(path)[1].lower()

        if extension == ".docx":
            try:
                from docx import Document
                doc = Document(path)
                return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            except Exception as e:
                print(f"[ERREUR DOCX] {path} -> {e}")
                return ""

        elif extension in [".xlsx", ".xls"]:
            try:
                df = pd.read_excel(path)
                return df.to_string(index=False)
            except Exception as e:
                print(f"[ERREUR EXCEL] {path} -> {e}")
                return ""

        elif extension == ".csv":
            try:
                df = pd.read_csv(path)
                return df.to_string(index=False)
            except Exception as e:
                print(f"[ERREUR CSV] {path} -> {e}")
                return ""

        elif extension == ".txt":
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception as e:
                print(f"[ERREUR TXT] {path} -> {e}")
                return ""

        else:
            print(f"[AVERTISSEMENT] Extension non supportée : {extension}")
            return ""



    # --- Construction du prompt ---
    def construire_prompt_competences(
        self,
        nom_emploi,
        contenu_fichier=None,
        params=None,
        competences_detectees=None,
        competences_projet=None,
        niveaux=None,
        infos_projet=None,
        emplois_existants=None
        ):
        """
        Construit dynamiquement le prompt pour générer les compétences d’un emploi.
        """
        import json
        params = params or {}
        competences_detectees = competences_detectees or []
        competences_projet = competences_projet or []
        niveaux = niveaux or []
        infos_projet = infos_projet or {}
        emplois_existants = emplois_existants or []

        prompt = f"""
        Tu es un expert RH. Ton objectif est de générer les compétences clés associées à un poste.

        Poste : "{nom_emploi}"
        {('Ceci est un poste de management.' if params.get('est_management') else '')}
        Pourcentage approximatif de soft skills souhaité : {params.get('pourcentage_soft', 30)}%
        """

        # 🔹 Contexte entreprise
        prompt += f"""
        Contexte de l’entreprise :
        - Secteur : {infos_projet.get('secteur_activite', 'Inconnu')}
        - Type : {infos_projet.get('type_entreprise', 'Inconnu')}
        - Taille : {infos_projet.get('taille', 'Inconnue')}
        - Nombre de filiales : {infos_projet.get('nb_filiales', 'Inconnu')}
        - International : {'Oui' if infos_projet.get('est_international') else 'Non'}
        - Niveau de digitalisation : {infos_projet.get('niveau_digitalisation', 'Inconnu')}
        """

        # 🔹 Liste des autres emplois (pour contexte)
        autres_emplois = [e['titre'] for e in emplois_existants if e['titre'] != nom_emploi]
        prompt += f"""
        Autres emplois du référentiel :
        {json.dumps(autres_emplois[:30], ensure_ascii=False, indent=2)}
        """

        # 🔹 Compétences projet (si dispo)
        if competences_projet:
            dictionnaire = {}
            for comp in competences_projet:
                comp_nom = comp["nom"]
                niveau_map = comp.get("descriptions", {})
                dictionnaire[comp_nom] = niveau_map

            prompt += f"""
            Tu dois t’appuyer sur ce dictionnaire de compétences existant pour harmoniser la terminologie, et tu peux proposer et enrichir au besoin:
            {json.dumps(dictionnaire, ensure_ascii=False, indent=2)}
            """

        # 🔹 Définitions des niveaux (si dispo)
        definitions_niveaux = {str(n.get("position")): n.get("definition") for n in niveaux if n.get("definition")}
        if definitions_niveaux:
            prompt += f"""
            Définitions générales des niveaux :
            {json.dumps(definitions_niveaux, ensure_ascii=False, indent=2)}
            """

        # 🔹 Selon le cas : contenu brut ou compétences extraites
        if competences_detectees:
            prompt += f"""
            Les compétences suivantes ont été extraites d’un fichier :
            {json.dumps(competences_detectees, indent=2)}
            Tu dois les utiliser comme base et les enrichir et structurer proprement.
            """
        elif contenu_fichier:
            prompt += f"""
            Voici un extrait de la description de poste :
            """
            {contenu_fichier[:1000]}
            """
            Utilise-le pour identifier les compétences attendues pour ce poste.
            """
        else:
            prompt += "Aucun document n’est fourni. Gère ce cas en te basant uniquement sur le titre du poste.\n"

        # 🔹 Types de compétences à produire
        prompt += """
        Types de compétences attendues :
        - Compétences cœur métier 
        - Compétences managériales 
        - Compétences transverses 
        """

        # 🔹 Paramètres supplémentaires
        prompt += f"""
        Je note que chaque compétence est déclinée en {infos_projet.get('nb_niveaux', 3)} niveaux.
        Génère entre {infos_projet.get('min_comp', 6)} et {infos_projet.get('max_comp', 12)} compétences.

        Retourne uniquement un JSON au format suivant :
        [
        {{ "nom": "Analyse de données", "type": "hard", "niveau_requis": 3 }},
        {{ "nom": "Communication", "type": "soft", "niveau_requis": 2 }}
        ]
        """

        return prompt

    def construire_prompt_competences2(self, titre_emploi, contenu_fichier, params):
        nb_niveaux = self.ui.spin_nombre_niveaux_competences.value()
        min_comp = self.ui.spin_min_nombre_competences.value()
        max_comp = self.ui.spin_max_nombre_competences.value()

        type_fichier = self.ui.combo_type_fichier.currentText()
        if type_fichier == "Autre":
            type_fichier = self.ui.input_autre_type.text().strip()

        cartographie = self.extraire_cartographie_depuis_tree() if self.ui.tree_cartographie_metiers.topLevelItemCount() > 0 else None

        prompt = f"Génère entre {min_comp} et {max_comp} compétences pour l'emploi suivant : {titre_emploi}.\n"
        prompt += f"Chaque compétence doit être déclinée en {nb_niveaux} niveaux.\n"
        prompt += f"Soft skills demandées : {params['pourcentage_soft']}%.\n"
        prompt += f"Poste de management : {'Oui' if params['est_management'] else 'Non'}.\n"

        if contenu_fichier:
            prompt += f"\nContenu du fichier ({type_fichier}) :\n{contenu_fichier[:1000]}...\n"

        if cartographie:
            prompt += f"\nCartographie des emplois à respecter :\n{json.dumps(cartographie, ensure_ascii=False)}\n"

        # TODO: Ajouter référentiel et dictionnaire si existants

        prompt += "\nRetourne uniquement un JSON de la forme : [{\"nom\": \"...\", \"niveau_requis\": 2, \"type\": \"hard\"}, ...]"
        return prompt


    def show_tree_context_menu(self, position):
        item = self.ui.tree_cartographie_metiers.itemAt(position)
        if not item:
            return

        menu = QMenu()
        ajouter = menu.addAction("Ajouter un sous-élément")
        renommer = menu.addAction("Renommer")
        supprimer = menu.addAction("Supprimer")

        action = menu.exec_(self.ui.tree_cartographie_metiers.viewport().mapToGlobal(position))

        if action == ajouter:
            self.ajouter_sous_element(item)
        elif action == renommer:
            self.ui.tree_cartographie_metiers.editItem(item)
        elif action == supprimer:
            self.supprimer_item(item)

    def ajouter_sous_element(self, parent_item):
        enfant = QTreeWidgetItem(["Nouvel élément"])
        enfant.setFlags(enfant.flags() | QtCore.Qt.ItemIsEditable)
        parent_item.addChild(enfant)
        parent_item.setExpanded(True)
        self.ui.tree_cartographie_metiers.setCurrentItem(enfant)
        self.ui.tree_cartographie_metiers.editItem(enfant)
    
    def supprimer_item(self, item):
        parent = item.parent()
        if parent:
            parent.removeChild(item)
        else:
            index = self.ui.tree_cartographie_metiers.indexOfTopLevelItem(item)
            self.ui.tree_cartographie_metiers.takeTopLevelItem(index)
    def edit_tree_item(self, item, column):
        self.ui.tree_cartographie_metiers.editItem(item, column)
    
    def ajouter_famille(self):
        item = QTreeWidgetItem(["Nouvelle famille"])
        item.setFlags(item.flags() | QtCore.Qt.ItemIsEditable)
        self.ui.tree_cartographie_metiers.addTopLevelItem(item)
        self.ui.tree_cartographie_metiers.setCurrentItem(item)
        self.ui.tree_cartographie_metiers.editItem(item)

    def ajouter_familles_par_liste(self):
        text, ok = QInputDialog.getMultiLineText(self, "Ajouter plusieurs familles", "Une par ligne :")
        if ok:
            for nom in text.split("\n"):
                if nom.strip():
                    item = QTreeWidgetItem([nom.strip()])
                    self.ui.tree_cartographie_metiers.addTopLevelItem(item)

    def ajouter_sous_familles_par_liste(self):
        parent = self.ui.tree_cartographie_metiers.currentItem()
        if not parent:
            QMessageBox.warning(self, "Erreur", "Sélectionnez une famille.")
            return

        text, ok = QInputDialog.getMultiLineText(self, "Ajouter sous_familles", "Une par ligne :")
        if ok:
            for nom in text.split("\n"):
                if nom.strip():
                    child = QTreeWidgetItem([nom.strip()])
                    parent.addChild(child)

    def ajouter_metiers_par_liste(self):
        parent = self.ui.tree_cartographie_metiers.currentItem()
        if not parent or not parent.parent():
            QMessageBox.warning(self, "Erreur", "Sélectionnez une sous-famille.")
            return

        text, ok = QInputDialog.getMultiLineText(self, "Ajouter métiers", "Une par ligne :")
        if ok:
            for nom in text.split("\n"):
                if nom.strip():
                    child = QTreeWidgetItem([nom.strip()])
                    parent.addChild(child)

    def charger_cartographie_depuis_base(self):
        if not self.job_controller:
            return

        data = self.job_controller.recuperer_hierarchie_as_json()
        print("[DEBUG carto chargée de la base]:", json.dumps(data, indent=2, ensure_ascii=False))  # 👈 Ajout
        if data:
            self.afficher_hierarchie_dans_tree(data)


    def importer_fichier_client(self):
        """Ouvre un sélecteur de fichier et extrait le contenu texte."""
        filepath, _ = QFileDialog.getOpenFileName(self, "Sélectionner un fichier", "", "Documents (*.docx *.xlsx *.xls)")
        
        if not filepath:
            return  # utilisateur a annulé

        extension = os.path.splitext(filepath)[1].lower()

        contenu = ""
        try:
            if extension == ".docx":
                doc = Document(filepath)
                contenu = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            elif extension in [".xlsx", ".xls"]:
                df = pd.read_excel(filepath)
                contenu = df.to_string(index=False)  # ou df.head(20).to_string() pour limiter
            else:
                QMessageBox.warning(self, "Fichier non supporté", "Seuls les fichiers .docx ou .xlsx sont supportés.")
                return

            self.contenu_fichier_client = contenu.strip()
            self.type_fichier_client = self.ui.input_type_fichier_importe.text().strip()
            QMessageBox.information(self, "Fichier chargé", "Contenu du fichier client importé avec succès.")

        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de l'import du fichier : {str(e)}")

    def sauvegarder_cartographie_actuelle(self):
        if not self.current_projet_id:
            QMessageBox.warning(self, "Projet requis", "Veuillez d’abord ouvrir un projet.")
            return

        data = self.extraire_cartographie_depuis_tree()

        try:
            self.job_controller.sauvegarder_hierarchie_depuis_json(data, self.current_projet_id)
            QMessageBox.information(self, "Succès", "Cartographie enregistrée avec succès.")
        except Exception as e:
            print("[Erreur sauvegarde finale] :", e)
            QMessageBox.critical(self, "Erreur", f"Échec de la sauvegarde : {str(e)}")

    
    def customize_layouts(self):
        #self.ui.input_definition_niveau_competence.setStyleSheet("background-color: lightblue;")

        #self.ui.verticalLayout_3.setStretch(0, 0)  # frame_15 = boutons du haut
        #self.ui.verticalLayout_3.setStretch(1, 1)  # stackedWidget = prend l'espace restant

        # === PAGE 1 : Cartographie des emplois (verticalLayout_9) ===
        # 0 = Titre
        # 1 = Tree widget
        # 2 = zone de chat IA
        # 3 = boutons "ajouter famille", etc.
        # 4 = bouton sauvegarder
        # 5 = bouton exporter

        self.ui.verticalLayout_7.setStretch(0, 0)  # frame_8
        self.ui.horizontalLayout_10.setStretch(1, 0)  # frame_10
        self.ui.horizontalLayout_10.setStretch(2, 1)  # input_definition_niveau_competence (stretch !)
        self.ui.horizontalLayout_10.setStretch(3, 0)  # bouton "ajouter compétence"

       
        self.ui.frame_38.setMaximumHeight(50)
        

    def extraire_cartographie_depuis_tree(self):
        """Parcourt le tree pour reconstruire le JSON actuel"""
        tree = self.ui.tree_cartographie_metiers
        familles = []

        for i in range(tree.topLevelItemCount()):
            item_fam = tree.topLevelItem(i)
            famille = {"nom": item_fam.text(0), "sous_familles": []}

            for j in range(item_fam.childCount()):
                item_sf = item_fam.child(j)
                sous_famille = {"nom": item_sf.text(0), "emplois": []}

                for k in range(item_sf.childCount()):
                    item_emp = item_sf.child(k)
                    sous_famille["emplois"].append(item_emp.text(0))

                famille["sous_familles"].append(sous_famille)

            familles.append(famille)

        return {"familles": familles}
    def discuter_avec_ia_cartographie(self):
        if not self.current_projet_id:
            QMessageBox.warning(self, "Projet requis", "Veuillez d’abord créer ou ouvrir un projet.")
            return

        message_user = self.ui.input_zone_tchat_carto_emplois.toPlainText().strip()
        if not message_user:
            QMessageBox.warning(self, "Message vide", "Veuillez écrire un message avant d’envoyer.")
            return

        # 🔸 1. Récupération du JSON actuel depuis le tree
        carto_actuelle = self.extraire_cartographie_depuis_tree()
        carto_json_str = json.dumps(carto_actuelle, indent=2, ensure_ascii=False)

        # 🔸 2. Contexte du projet
        secteur = self.ui.input_secteur_activite.text().strip()
        type_entreprise = self.ui.combo_type_entreprise.currentText().strip()
        taille = self.ui.input_taille_entreprise.text().strip()
        nb_filiales = self.ui.input_nb_filiales.text().strip()
        international = "internationale" if self.ui.check_internationale.isChecked() else "nationale"
        digital = self.ui.spin_niveau_digitalisation.value()

        # 🔸 3. Fichier client (texte déjà extrait)
        texte_fichier = self.contenu_fichier_client.strip()
        type_fichier = self.type_fichier_client.strip()

        # 🔸 4. Prompt complet
        prompt = (
            f"L'utilisateur souhaite affiner la cartographie des emplois suivante :\n{carto_json_str}\n\n"
            f"Voici le message utilisateur : {message_user}\n\n"
            f"Informations sur le contexte du projet :\n"
            f"- Secteur d'activité : {secteur}\n"
            f"- Type d'entreprise : {type_entreprise}\n"
            f"- Taille : {taille}\n"
            f"- Nombre de filiales : {nb_filiales}\n"
            f"- Portée : {international}\n"
            f"- Niveau de digitalisation : {digital}/10\n\n"
        )

        if texte_fichier:
            prompt += f"Contenu du fichier fourni ({type_fichier}) :\n{texte_fichier[:1000]}...\n\n"

        prompt += (
            "Retourne UNIQUEMENT le nouveau JSON structuré au format strict suivant :\n"
            '{ "familles": [ { "nom": "string", "sous_familles": [ { "nom": "string", "emplois": ["string", "string"] } ] } ] }.\n'
            " N’inclus AUCUN texte explicatif ou markdown. Réponds uniquement avec du JSON brut.\n"
        )

        print("[Prompt IA affinement] :", prompt)

        # 🔹 Appel à OpenAI
        try:
            reponse = self.openai_controller.openai_service.ask_once(prompt)
            print("[Réponse IA affinement] :", reponse)

            json_str = self.extraire_json_depuis_texte(reponse)
            data = json.loads(json_str)

            self.afficher_hierarchie_dans_tree(data)
            # 🔻 Vide le champ de saisie après envoi
            self.ui.input_zone_tchat_carto_emplois.clear()
            QMessageBox.information(self, "Cartographie mise à jour", "La cartographie a été affinée et rechargée dans l’arbre.")

        except Exception as e:
            print("[Erreur IA] :", e)
            QMessageBox.critical(self, "Erreur", f"Erreur lors de la réponse IA : {str(e)}")


    def generer_cartographie_via_ia(self):
        if not self.current_projet_id:
            QMessageBox.warning(self, "Projet requis", "Veuillez d’abord créer ou ouvrir un projet.")
            return

        # Récupère les paramètres du formulaire avec les bons noms
        secteur = self.ui.input_secteur_activite.text().strip()
        type_entreprise = self.ui.combo_type_entreprise.currentText().strip()
        taille = self.ui.input_taille_entreprise.text().strip()
        nb_filiales = self.ui.input_nb_filiales.text().strip()
        international = self.ui.check_internationale.isChecked()
        digital = self.ui.spin_niveau_digitalisation.value()
        commentaire = self.ui.input_commentaire_projet.toPlainText().strip()

        # Construction dynamique du prompt
        prompt = f"Génère une cartographie des emplois pour une entreprise dans le secteur '{secteur}'. "
        if type_entreprise:
            prompt += f"L'entreprise est de type '{type_entreprise}'. "
        if taille:
            prompt += f"Taille de l’entreprise : {taille}. "
        if international:
            prompt += f"C’est une entreprise internationale avec {nb_filiales} filiales. "
        else:
            prompt += f"C’est une entreprise nationale avec {nb_filiales} filiales. "
        prompt += f"Niveau de digitalisation : {digital}/10. "
        if self.contenu_fichier_client:
            prompt += f"\n\n📎 Un fichier client a été fourni (type : {self.type_fichier_client}). Voici son contenu :\n{self.contenu_fichier_client}\n"
            prompt += "Merci d’en tenir compte et l'utiliser comme source supplémentaire générer la cartographie des emplois."

        if commentaire:
            prompt += f"Contexte additionnel : {commentaire}. "

        prompt += 'Retourne un JSON structuré STRICTEMENT au format suivant : { "familles": [ { "nom": "string", "sous_familles": [{ "nom": "string", "emplois": ["string", "string"] } ]}] } Utilise EXCLUSIVEMENT les clés suivantes : familles , sous_familles , emplois'
       
        # Appel à OpenAI
        print("[Prompt généré] ", prompt)
        if len(prompt) > 15000:  # seuil de sécurité (~4000 tokens)
            QMessageBox.warning(self, "Fichier trop volumineux", 
                "Le contenu du fichier client est trop volumineux pour être traité.\n"
                "Merci de réduire sa taille (par exemple, garder uniquement les sections pertinentes).")
            return

        reponse = self.openai_controller.generer_cartographie(prompt)
        print("[Réponse IA brute] :")
        print(reponse)
        # 🧼 Nettoyage
        json_str = self.extraire_json_depuis_texte(reponse)
        try:
            data = json.loads(json_str)
            self.afficher_hierarchie_dans_tree(data)

            # Sauvegarde en base
            #####self.job_controller.sauvegarder_hierarchie_depuis_json(data, self.current_projet_id)

            QMessageBox.information(self, "Succès", "La cartographie a été générée.")
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de l’analyse de la réponse : {str(e)}")
            print("[Erreur parsing/sauvegarde] :", e)


    def afficher_hierarchie_dans_tree(self, data):
        #self.ui.tree_cartographie_metiers.clear()
        self.ui.tree_cartographie_metiers.clear()
        self.ui.tree_cartographie_metier_generateur_competence.clear()
        self.ui.tree_carto_emploi_fichedeposte.clear()
        familles = data.get("familles", [])
        print("[DEBUG - JSON à afficher dans le tree]")
        
        
        # Icônes personnalisées
        icone_famille = QtGui.QIcon(":/icon/folder.ico")  # à adapter selon ton chemin/ressource
        icone_sf = QtGui.QIcon(":/icon/folder-git-2.ico")
        icone_emploi = QtGui.QIcon(":/icon/briefcase.ico")
        for famille in familles:
            nom_famille = famille.get("nom", "Sans nom")
            item_famille = QTreeWidgetItem([nom_famille])
            item_famille.setIcon(0, icone_famille)  # <- 🟢 Ajout icône
            item_famille.setFlags(item_famille.flags() | QtCore.Qt.ItemIsEditable)
            self.ui.tree_cartographie_metiers.addTopLevelItem(item_famille)

            sous_familles = famille.get("sous_familles", [])  # ✅
            for sous_famille in sous_familles:
                nom_sf = sous_famille.get("nom", "Sans nom")
                item_sf = QTreeWidgetItem([nom_sf])
                item_sf.setIcon(0, icone_sf)  # <- 🟢 Ajout icône
                item_sf.setFlags(item_sf.flags() | QtCore.Qt.ItemIsEditable)
                item_famille.addChild(item_sf)

                emplois = sous_famille.get("emplois", [])
                for emploi in emplois:
                    item_emploi = QTreeWidgetItem([emploi])
                    item_emploi.setIcon(0, icone_emploi)  # <- 🟢 Ajout icône
                    item_emploi.setFlags(item_emploi.flags() | QtCore.Qt.ItemIsEditable)
                    item_sf.addChild(item_emploi)

        self.ui.tree_cartographie_metiers.expandAll()

    def sauvegarder_parametres_projet(self):
        if not self.current_projet_id:
            QMessageBox.warning(self, "Projet non chargé", "Veuillez d’abord créer ou ouvrir un projet.")
            return

        # Récupération des valeurs modifiées depuis l’UI
        data = {
            "secteur": self.ui.input_secteur_activite.text().strip(),
            "type_entreprise": self.ui.combo_type_entreprise.currentText(),
            "taille": self.ui.input_taille_entreprise.text().strip(),
            "filiales": int(self.ui.input_nb_filiales.text() or 0),
            "digital": int(self.ui.spin_niveau_digitalisation.value()),
            "est_international": self.ui.check_internationale.isChecked(),
            "nb_niveaux": self.ui.spin_nombre_niveaux_competences.value(),
            "nb_comp_min": self.ui.spin_min_nombre_competences.value(),
            "nb_comp_max": self.ui.spin_max_nombre_competences.value(),
            #"nb_macro_min": self.ui.spin_min_macro.value(),
            #"nb_macro_max": self.ui.spin_max_macro.value(),
            #"nb_micro_min": self.ui.spin_min_micro.value(),
            #"nb_micro_max": self.ui.spin_max_micro.value(),
            "nb_niveaux_dict": self.ui.spin_nombre_niveaux_competences_dictionnaire.value()
        }

        # Mise à jour en base via le CRUD
        # 1. On récupère les données actuelles du projet
        projet = crud.get_projet_by_id(self.current_projet_id)  # Il te faut cette fonction si tu ne l’as pas encore

        # 2. On complète les champs manquants avec les valeurs existantes
        crud.update_projet_parametres(
            self.current_projet_id,
            secteur_activite=data["secteur"],
            type_entreprise=data["type_entreprise"],
            taille=data["taille"],
            est_international=data["est_international"],
            nb_filiales=data["filiales"],
            niveau_digitalisation=data["digital"],
            nb_niveaux=data["nb_niveaux"],
            nb_competences_min=data["nb_comp_min"],
            nb_competences_max=data["nb_comp_max"],
            nb_macro_activites_min=projet[10],  # index selon l’ordre de tes colonnes
            nb_macro_activites_max=projet[11],
            nb_micro_activites_min=projet[12],
            nb_micro_activites_max=projet[13]
        )


        QMessageBox.information(self, "Succès", "Les paramètres du projet ont bien été mis à jour ✅")
  
        # Recharge les données mises à jour dans les champs
        self.charger_infos_projet_dans_formulaire(crud.get_projet_by_id(self.current_projet_id))
        self.set_inputs_projet_enabled(False)

    
    def set_inputs_projet_enabled(self, enabled: bool):
        self.ui.input_secteur_activite.setEnabled(enabled)
        self.ui.combo_type_entreprise.setEnabled(enabled)
        self.ui.input_taille_entreprise.setEnabled(enabled)
        self.ui.input_nb_filiales.setEnabled(enabled)
        self.ui.spin_niveau_digitalisation.setEnabled(enabled)
        self.ui.check_internationale.setEnabled(enabled)
        self.ui.btn_sauvegarder_inputs_projet.setEnabled(enabled)


    def _verrouiller_pages(self):
        self.ui.page_cartographie.setEnabled(False)
        self.ui.page_competences_poste.setEnabled(False)
        self.ui.page_referentiel.setEnabled(False)
        self.ui.page_dictionnaire.setEnabled(False)

    def _activer_pages(self):
        self.ui.page_cartographie.setEnabled(True)
        self.ui.page_competences_poste.setEnabled(True)
        self.ui.page_referentiel.setEnabled(True)
        self.ui.page_dictionnaire.setEnabled(True)
        
    def charger_projet(self, projet_id):
        from db import crud
        self.ui.tree_cartographie_metiers.clear()
        self.ui.tree_cartographie_metier_generateur_competence.clear()
        self.ui.tree_carto_emploi_fichedeposte.clear()
        self.current_projet_id = projet_id
        projet_data = crud.get_projet_by_id(self.current_projet_id)

        self.job_controller = JobController(self.current_projet_id)
        self.competence_controller = CompetenceController(self.current_projet_id)
        self.finalite_controller = FinaliteEmploiController(self.current_projet_id)
        self.fiche_poste_controller = FichePosteController(self.current_projet_id)
        self.openai_controller = OpenAIController()

        self.charger_infos_projet_dans_formulaire(projet_data)
        self.charger_cartographie_depuis_base()

        if self.job_controller:
            data = self.job_controller.recuperer_hierarchie_as_json()
            if data:
                self.afficher_hierarchie_dans_tree_generateur(data)

        self.update_combo_sous_famille()

        competences = self.competence_controller.charger_referentiel()
        self.afficher_referentiel_dans_liste(competences)

        self.charger_dictionnaire_depuis_base()
        self.charger_infos_projet_dans_formulaire(projet_data)
        self._activer_pages()
        self.clear_layout(self.ui.layout_zone_fichiers)
        self.ui.btn_ajouter_zone_fichier.setEnabled(True)

        

   

    def creer_nouveau_projet(self):
        dialog = DialogNouveauProjet()
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()

            if not data["nom"] or not data["secteur"]:
                QMessageBox.warning(self, "Champs requis", "Le nom du projet et le secteur sont obligatoires.")
                return

            crud.create_projet(
                data["nom"], data["secteur"], data["type_entreprise"], data["taille"],
                data["international"], data["filiales"],
                data["digitalisation"], data["nb_niveaux"], data["nb_comp_min"], data["nb_comp_max"],
                data["nb_macro_min"], data["nb_macro_max"], data["nb_micro_min"], data["nb_micro_max"]
            )

            projets = crud.get_all_projets()
            self.current_projet_id = projets[-1][0]

            # Création des niveaux par défaut
            nb_niveaux = data["nb_niveaux"]
            for i in range(1, nb_niveaux + 1):
                libelle = f"Niveau {i}"
                crud.add_niveau(self.current_projet_id, libelle, i)

            self.charger_projet(self.current_projet_id)

            QMessageBox.information(self, "Projet créé", f"Projet « {data['nom']} » créé avec succès.")

    def ouvrir_projet(self):
        from db import crud

        projets = crud.get_all_projets()
        if not projets:
            QMessageBox.warning(self, "Aucun projet", "Aucun projet n’est disponible.")
            return

        items = [f"{p[0]} – {p[1]}" for p in projets]  # id – nom
        item, ok = QInputDialog.getItem(self, "Ouvrir un projet", "Sélectionnez :", items, 0, False)
        
        if ok and item:
            projet_id = int(item.split("–")[0].strip())
            self.charger_projet(projet_id)



    def charger_infos_projet_dans_formulaire(self, projet):
        
        self.ui.input_secteur_activite.setText(projet["secteur_activite"])
        self.ui.combo_type_entreprise.setCurrentText(projet["type_entreprise"] or "")
        self.ui.input_taille_entreprise.setText(projet["taille"] or "")
        self.ui.input_nb_filiales.setText(str(projet["nb_filiales"] or ""))
        self.ui.spin_niveau_digitalisation.setValue(projet["niveau_digitalisation"] or 0)
        self.ui.check_internationale.setChecked(bool(projet["est_international"]))
        self.ui.spin_nombre_niveaux_competences.setValue(projet["nb_niveaux"] or 0)
        self.ui.spin_min_nombre_competences.setValue(projet["nb_competences_min"] or 0)
        self.ui.spin_max_nombre_competences.setValue(projet["nb_competences_max"] or 0)
        self.ui.spin_nombre_niveaux_competences_dictionnaire.setValue(projet["nb_niveaux"] or 0)
        self.ui.btn_ajouter_zone_fichier.setEnabled(True)
        self.ui.spin_min_macro_activites.setValue(projet["nb_macro_activites_min"] or 0)
        self.ui.spin_max_macro_activites.setValue(projet["nb_macro_activites_max"] or 0)
        self.ui.spin_min_micro_activites.setValue(projet["nb_micro_activites_min"] or 0)
        self.ui.spin_max_micro_activites.setValue(projet["nb_micro_activites_max"] or 0)
        
        self.ui.combo_niveau_requis.clear()
        self.ui.combo_niveau_competence_a_definir.clear()
        for i in range(1, projet["nb_niveaux"] + 1):
            self.ui.combo_niveau_competence_a_definir.addItem(str(i))
            self.ui.combo_niveau_requis.addItem(str(i))
        self.set_inputs_projet_enabled(False)
        self.set_inputs_projet_enabled(False)

    def afficher_hierarchie_dans_tree_generateur(self, data):
        from PyQt5.QtGui import QBrush, QColor

        tree = self.ui.tree_cartographie_metier_generateur_competence
        tree.clear()

        icone_famille = QtGui.QIcon(":/icon/folder.ico")
        icone_sf = QtGui.QIcon(":/icon/folder-git-2.ico")
        icone_emploi = QtGui.QIcon(":/icon/briefcase.ico")

        # 📦 Emplois traités (ayant des compétences)
        emplois_traite_ids = self.competence_controller.get_emploi_ids_avec_competences()
        emplois_existants = self.job_controller.get_liste_emplois()
        emploi_nom_par_id = {emp["id"]: emp["titre"] for emp in emplois_existants}
        emploi_id_par_nom = {emp["titre"]: emp["id"] for emp in emplois_existants}

        for famille in data.get("familles", []):
            item_fam = QTreeWidgetItem([famille["nom"]])
            item_fam.setIcon(0, icone_famille)
            tree.addTopLevelItem(item_fam)

            for sous_famille in famille.get("sous_familles", []):
                item_sf = QTreeWidgetItem([sous_famille["nom"]])
                item_sf.setIcon(0, icone_sf)
                item_fam.addChild(item_sf)

                for emploi_nom in sous_famille.get("emplois", []):
                    item_emp = QTreeWidgetItem([emploi_nom])
                    item_emp.setIcon(0, icone_emploi)

                    emploi_id = emploi_id_par_nom.get(emploi_nom)
                    if emploi_id in emplois_traite_ids:
                        item_emp.setForeground(0, QBrush(QColor("green")))
                        item_emp.setIcon(0, QIcon(":/icon/emploi-selectionne2_resized.ico"))
                    else:
                        item_emp.setForeground(0, QBrush(QColor("black")))

                    item_sf.addChild(item_emp)

        tree.expandAll()
        
    def afficher_hierarchie_dans_tree_generateur_v1(self, data):
        tree = self.ui.tree_cartographie_metier_generateur_competence
        tree.clear()
        # Icônes personnalisées
        icone_famille = QtGui.QIcon(":/icon/folder.ico")  # à adapter selon ton chemin/ressource
        icone_sf = QtGui.QIcon(":/icon/folder-git-2.ico")
        icone_emploi = QtGui.QIcon(":/icon/briefcase.ico")
        for famille in data.get("familles", []):
            item_fam = QTreeWidgetItem([famille["nom"]])
            item_fam.setIcon(0, icone_famille)
            tree.addTopLevelItem(item_fam)

            for sous_famille in famille.get("sous_familles", []):
                item_sf = QTreeWidgetItem([sous_famille["nom"]])
                item_sf.setIcon(0, icone_sf)  # <- 🟢 Ajout icône
                item_fam.addChild(item_sf)

                for emploi in sous_famille.get("emplois", []):
                    item_emp = QTreeWidgetItem([emploi])
                    item_emp.setIcon(0, icone_emploi)
                    item_sf.addChild(item_emp)

        tree.expandAll()

    def afficher_hierarchie_dans_tree_fiche_poste(self, data):
        tree = self.ui.tree_carto_emploi_fichedeposte
        tree.clear()
        # Icônes personnalisées
        icone_famille = QtGui.QIcon(":/icon/folder.ico")  # à adapter selon ton chemin/ressource
        icone_sf = QtGui.QIcon(":/icon/folder-git-2.ico")
        icone_emploi = QtGui.QIcon(":/icon/briefcase.ico")
        for famille in data.get("familles", []):
            item_fam = QTreeWidgetItem([famille["nom"]])
            item_fam.setIcon(0, icone_famille)
            tree.addTopLevelItem(item_fam)

            for sous_famille in famille.get("sous_familles", []):
                item_sf = QTreeWidgetItem([sous_famille["nom"]])
                item_sf.setIcon(0, icone_sf)  # <- 🟢 Ajout icône
                item_fam.addChild(item_sf)

                for emploi in sous_famille.get("emplois", []):
                    item_emp = QTreeWidgetItem([emploi])
                    item_emp.setIcon(0, icone_emploi)
                    item_sf.addChild(item_emp)

        tree.expandAll()


    def on_cartographie_emploi_toggled(self):
        #self.display_achats_in_table()
        self.ui.stackedWidget.setCurrentIndex(0)

    def on_generateur_competences_emploi_toggled(self):
        #self.display_achats_in_table()
        self.ui.stackedWidget.setCurrentIndex(1)
        # Recharger la cartographie à chaque fois qu'on revient sur l'onglet
        if self.job_controller:
            data = self.job_controller.recuperer_hierarchie_as_json()
            if data:
                self.afficher_hierarchie_dans_tree_generateur(data)

    def on_referentiel_toggled(self):
        #self.display_achats_in_table()
        self.ui.stackedWidget.setCurrentIndex(2)

    def on_dictionnaire_toggled(self):
        #self.display_achats_in_table()
        self.ui.stackedWidget.setCurrentIndex(3)

    def on_carto_niveau_competence_toggled(self):
        #self.display_achats_in_table()
        self.ui.stackedWidget.setCurrentIndex(4)

    def on_fiche_poste_toggled(self):
        #self.display_achats_in_table()
        self.ui.stackedWidget.setCurrentIndex(5)
        if self.job_controller:
            data = self.job_controller.recuperer_hierarchie_as_json()
            if data:
                self.afficher_hierarchie_dans_tree_fiche_poste(data)

    def on_parametres_toggled(self):
        #self.display_achats_in_table()
        self.ui.stackedWidget.setCurrentIndex(6)

if __name__ == "__main__":
    init_database()  # Initialiser la BDD au lancement
    app = QApplication(sys.argv)
    window = MainApp()
    sys.exit(app.exec_())
