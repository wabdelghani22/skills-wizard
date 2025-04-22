# export_fiches.py

import os
import sqlite3
from docx import Document
from PyQt5.QtWidgets import QMessageBox

DB_PATH = os.path.join(os.path.dirname(__file__), "db/competence_app.db")
# Construction du chemin vers la base de données
#DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "competence_app.db")

def get_emplois_remplis():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("""
        SELECT emplois.id, emplois.titre, emplois.finalite
        FROM emplois
        WHERE emplois.finalite IS NOT NULL AND emplois.titre IS NOT NULL
    """)

    emplois = cursor.fetchall()
    conn.close()
    return emplois

def get_competences_by_emploi(emploi_id):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("""
        SELECT competences.nom
        FROM competences
        JOIN emploi_competence ON competences.id = emploi_competence.competence_id
        WHERE emploi_competence.emploi_id = ?
    """, (emploi_id,))

    competences = [row[0] for row in cursor.fetchall()]
    conn.close()
    return competences

def get_activites_by_emploi(emploi_id):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("""
        SELECT macro_activites.contenu AS macro_activite, micro_activites.contenu AS micro_activite
        FROM macro_activites
        LEFT JOIN micro_activites ON macro_activites.id = micro_activites.macro_id
        WHERE macro_activites.emploi_id = ?
    """, (emploi_id,))

    activites = cursor.fetchall()
    conn.close()
    return activites

def create_export_directory():
    export_dir = os.path.join(os.path.dirname(__file__), "export_fiches")
    if not os.path.exists(export_dir):
        os.makedirs(export_dir)
    return export_dir

def create_docx_file(emploi, competences, activites, export_dir):
    doc = Document()

    doc.add_heading(emploi[1], 0)  # Titre de la fiche de poste

    doc.add_heading('Finalité du poste:', level=1)
    doc.add_paragraph(emploi[2])  # Finalité

    doc.add_heading('Compétences requises:', level=1)
    for competence in competences:
        doc.add_paragraph(f"- {competence}")

    doc.add_heading('Activités:', level=1)
    for activite in activites:
        macro_activite = activite[0]
        micro_activite = activite[1]
        doc.add_paragraph(f"Macro-activité: {macro_activite}")
        if micro_activite:
            doc.add_paragraph(f"Micro-activité: {micro_activite}", style='List Bullet')

    file_path = os.path.join(export_dir, f"{emploi[1]}.docx")
    doc.save(file_path)

def export_fiches(self):
    emplois = get_emplois_remplis()

    if not emplois:
        QMessageBox.warning(self, "Aucun emploi", "Aucun emploi à exporter.")
        return

    export_dir = create_export_directory()

    for emploi in emplois:
        emploi_id = emploi[0]
        competences = get_competences_by_emploi(emploi_id)
        activites = get_activites_by_emploi(emploi_id)

        create_docx_file(emploi, competences, activites, export_dir)

    QMessageBox.information(self, "Succès", "Les fiches de poste ont été exportées avec succès.")
